# -*- coding: utf-8 -*-
import sys
from pathlib import Path
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

TEMPLATE = Path(r"C:\Users\ADMIN\Downloads\BaoCao_chuong3_4_5nhom9.docx")

doc = Document(TEMPLATE)

print("=" * 80)
print("PARAGRAPHS:")
print("=" * 80)
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else "Normal"
    text = p.text.strip()
    if text:
        print(f"[{i}] ({style}) {text[:200]}")

print()
print("=" * 80)
print("TABLES:")
print("=" * 80)
for ti, t in enumerate(doc.tables):
    print(f"--- Table {ti} ({len(t.rows)} rows x {len(t.columns)} cols) ---")
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip().replace("\n", " | ")
            if txt:
                print(f"  [{ri},{ci}] {txt[:200]}")
