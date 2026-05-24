# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document


TEMPLATE = Path(r"C:\Users\ADMIN\Downloads\BaoCaoTuan6.docx")
OUTPUT_DIR = Path(r"C:\Users\ADMIN\Downloads")


REPORTS = {
    7: [
        ("1. Nội dung và Mục tiêu", [
            "Tiếp tục cập nhật tiến độ sau giai đoạn đã xây dựng được server vulnerable, server secured và mô hình Liveness Detection ban đầu.",
            "Rà soát lại quy trình sinh trắc học theo đúng thứ tự: thu nhận ảnh, tiền xử lý, kiểm tra sống, trích chọn đặc trưng, so khớp và quyết định xác thực.",
            "Bổ sung dữ liệu thật/giả để cải thiện độ ổn định của mô hình trong các trường hợp ánh sáng yếu, ảnh chụp từ màn hình, ảnh in và ảnh chụp bằng điện thoại.",
            "Tối ưu bước tiền xử lý trước khi đưa ảnh vào mô hình, nhằm giảm lỗi do khuôn mặt bị lệch, thiếu sáng hoặc crop chưa đúng vùng mặt.",
        ]),
        ("2. Công việc đã thực hiện", [
            "Thu thập thêm mẫu khuôn mặt thật từ webcam trong nhiều điều kiện khác nhau, bao gồm chính diện, hơi nghiêng mặt, ánh sáng phòng và ánh sáng yếu.",
            "Tạo thêm mẫu giả mạo bằng ảnh hiển thị trên điện thoại, ảnh hiển thị trên màn hình máy tính và ảnh in giấy để mô phỏng các dạng spoofing phổ biến.",
            "Chuẩn hóa ảnh đầu vào: phát hiện khuôn mặt, crop vùng mặt, resize về cùng kích thước, cân bằng sáng và loại bỏ các ảnh bị mờ hoặc không nhận diện được mặt.",
            "Kiểm tra lại luồng xử lý trên server secured để bảo đảm mọi request xác thực đều phải đi qua bước Liveness Check trước khi thực hiện so khớp khuôn mặt.",
            "Ghi chú lại các bước xử lý theo quy trình biometric để đưa vào báo cáo, tránh trình bày lẫn lộn giữa phần tấn công và phần phòng vệ.",
        ]),
        ("3. Kết quả đạt được", [
            "Bộ dữ liệu đã đầy đủ và cân bằng hơn giữa nhóm Real và Fake, giúp quá trình huấn luyện/tinh chỉnh mô hình có cơ sở tốt hơn so với tuần trước.",
            "Pipeline tiền xử lý ổn định hơn, giảm tình trạng ảnh bị lệch mặt, thiếu sáng hoặc sai kích thước làm ảnh hưởng đến kết quả dự đoán của mô hình.",
            "Hệ thống demo đã vận hành đúng thứ tự xử lý: nhận ảnh từ client/API, tiền xử lý, kiểm tra sống, sau đó mới trích đặc trưng và so khớp danh tính.",
            "Đã cập nhật rõ tình hình hiện tại của đồ án: phần tấn công đã hoàn thành, phần phòng vệ đã tích hợp, đang chuyển sang tối ưu và đánh giá lại kết quả.",
            "Tiến độ chung đạt khoảng 80%, các chức năng chính của demo đã chạy được và bắt đầu tập trung vào cải thiện độ tin cậy.",
        ]),
        ("4. Khó khăn", [
            "Một số ảnh giả mạo chất lượng cao vẫn dễ gây nhầm lẫn nếu ánh sáng và góc chụp gần giống ảnh thật.",
            "Dữ liệu tự thu thập còn giới hạn về số người, thiết bị chụp và môi trường thực nghiệm nên khả năng tổng quát của mô hình chưa thật sự cao.",
            "Việc tiền xử lý cần cân bằng: xử lý quá mạnh có thể làm mất đặc trưng ảnh, còn xử lý quá nhẹ thì mô hình dễ bị ảnh hưởng bởi nhiễu và ánh sáng.",
        ]),
        ("5. Kế hoạch tiếp theo", [
            "Huấn luyện/tinh chỉnh lại mô hình Liveness Detection với bộ dữ liệu đã bổ sung và pipeline tiền xử lý ổn định hơn.",
            "Đánh giá lại các chỉ số FAR, FRR, APCER sau khi cập nhật dữ liệu, từ đó so sánh với kết quả tuần 6.",
            "Tiếp tục hoàn thiện phần báo cáo thực nghiệm, đặc biệt là mô tả quy trình biometric và vị trí tấn công/phòng vệ trong từng bước.",
            "Chuẩn bị kịch bản demo rõ ràng để tuần sau kiểm thử lại toàn bộ luồng từ tấn công đến phòng vệ.",
        ]),
    ],
    8: [
        ("1. Nội dung và Mục tiêu", [
            "Tập trung vào các bước sau trong quy trình sinh trắc học: trích chọn đặc trưng, so khớp khuôn mặt và quyết định xác thực.",
            "Tinh chỉnh ngưỡng nhận diện khuôn mặt và ngưỡng Liveness để cân bằng giữa bảo mật và khả năng sử dụng của hệ thống.",
            "Kiểm thử lại server vulnerable và server secured bằng cùng một nhóm dữ liệu để thấy rõ khác biệt trước và sau khi có cơ chế phòng vệ.",
            "Hoàn thiện kịch bản demo theo luồng hoàn chỉnh: đăng ký/nhận mẫu, xác thực thật, tấn công giả mạo, sau đó bật phòng vệ và kiểm tra lại.",
        ]),
        ("2. Công việc đã thực hiện", [
            "Huấn luyện/tinh chỉnh lại mô hình Liveness Detection sau khi bổ sung dữ liệu ở tuần 7, giữ kiến trúc MobileNetV2 nhưng điều chỉnh lại ngưỡng phân loại.",
            "Kiểm thử nhiều trường hợp: người thật đăng nhập bình thường, ảnh in, ảnh hiển thị trên điện thoại, ảnh chụp màn hình và request injection trực tiếp qua API.",
            "Ghi nhận kết quả từng bước trong quá trình xác thực, gồm điểm Liveness, kết quả phát hiện khuôn mặt, kết quả so khớp và phản hồi cuối cùng của server.",
            "So sánh server vulnerable và server secured: server vulnerable vẫn nhận request ảnh giả nếu dữ liệu khớp định dạng, còn server secured chặn trước ở bước kiểm tra sống.",
            "Chỉnh sửa nội dung báo cáo phần thiết kế hệ thống để thể hiện rõ hai nhánh: luồng tấn công và luồng phòng vệ.",
        ]),
        ("3. Kết quả đạt được", [
            "Tỷ lệ chặn ảnh giả mạo được cải thiện so với bản mô hình ban đầu, đặc biệt với ảnh in giấy và ảnh hiển thị trên điện thoại.",
            "Các request injection tiếp tục bị server secured từ chối trước khi bước so khớp danh tính diễn ra, cho thấy Liveness Detection đang được đặt đúng vị trí trong luồng biometric.",
            "Luồng xác thực đã rõ ràng: nhận ảnh -> tiền xử lý -> kiểm tra sống -> trích đặc trưng -> so khớp -> quyết định cho phép hoặc từ chối.",
            "Đã có đủ dữ liệu để trình bày bảng so sánh trước/sau, gồm các trường hợp đăng nhập hợp lệ, spoofing bằng ảnh và injection qua API.",
            "Tiến độ chung đạt khoảng 90%, phần demo và phần thực nghiệm gần hoàn chỉnh, chỉ còn chạy đánh giá cuối và tổng hợp trình bày.",
        ]),
        ("4. Khó khăn", [
            "Việc chọn ngưỡng quá chặt có thể làm tăng FRR, khiến người thật bị từ chối trong một số ảnh thiếu sáng hoặc góc mặt chưa chuẩn.",
            "Nếu chọn ngưỡng quá dễ, hệ thống có thể giảm FRR nhưng lại làm tăng nguy cơ ảnh giả mạo được chấp nhận.",
            "Một số số liệu cần chạy lại nhiều lần vì tập test còn nhỏ, nếu chỉ chạy một lần thì kết quả có thể chưa phản ánh đúng độ ổn định của hệ thống.",
        ]),
        ("5. Kế hoạch tiếp theo", [
            "Chạy đánh giá cuối cùng và tổng hợp bảng kết quả trước/sau khi áp dụng Liveness Detection.",
            "Quay video demo hoàn chỉnh: tấn công thành công trên server vulnerable và bị chặn trên server secured.",
            "Hoàn thiện chương kết quả thực nghiệm, bổ sung nhận xét về FAR, FRR, APCER và giới hạn của mô hình hiện tại.",
            "Kiểm tra lại file code, dữ liệu mẫu và slide để bảo đảm nội dung trình bày thống nhất với demo.",
        ]),
    ],
    9: [
        ("1. Nội dung và Mục tiêu", [
            "Hoàn thiện giai đoạn đánh giá, tổng hợp kết quả và chuẩn bị báo cáo cuối cho đề tài nghiên cứu tấn công vào hệ thống sinh trắc học.",
            "Kiểm tra lại toàn bộ thứ tự các bước biometric: enrollment, thu nhận mẫu, tiền xử lý, kiểm tra sống, trích đặc trưng, so khớp, quyết định xác thực và ghi nhận kết quả.",
            "Đảm bảo báo cáo thể hiện rõ hệ thống ban đầu có lỗ hổng ở đâu, kịch bản tấn công khai thác như thế nào và giải pháp Liveness Detection giảm rủi ro ra sao.",
            "Chuẩn bị tài liệu, slide và video demo phục vụ trình bày đồ án, ưu tiên nội dung ngắn gọn, dễ theo dõi và đúng kết quả thực nghiệm.",
        ]),
        ("2. Công việc đã thực hiện", [
            "Chạy lại các kịch bản kiểm thử cuối cùng trên server vulnerable và server secured để xác nhận kết quả trước khi đưa vào báo cáo.",
            "Tổng hợp số liệu FAR, FRR, APCER và so sánh trước/sau khi áp dụng Liveness Detection theo cùng một cách đo để tránh sai lệch.",
            "Rà soát lại code demo gồm client/server, script tấn công injection/replay, mô hình Liveness và dữ liệu mẫu dùng trong phần trình bày.",
            "Cập nhật nội dung báo cáo theo đúng mạch: cơ sở lý thuyết, mô hình hệ thống, điểm tấn công, triển khai demo, phòng vệ, đánh giá và kết luận.",
            "Chuẩn bị kịch bản thuyết trình: giới thiệu quy trình biometric, chỉ ra điểm yếu, chạy tấn công, sau đó chạy server secured để chứng minh biện pháp phòng vệ.",
        ]),
        ("3. Kết quả đạt được", [
            "Hệ thống vulnerable vẫn bị bypass bằng injection/replay trong môi trường lab, cho thấy nếu server chỉ tin dữ liệu ảnh gửi lên thì nguy cơ giả mạo là rất rõ ràng.",
            "Server secured chặn được phần lớn mẫu giả mạo nhờ bước Liveness Detection được đặt trước bước so khớp khuôn mặt.",
            "Báo cáo đã có đủ mạch nội dung: lý thuyết sinh trắc học, các bề mặt tấn công, kịch bản khai thác, giải pháp phòng vệ và kết quả đánh giá định lượng.",
            "Video demo và slide đã có nội dung chính để trình bày, tập trung vào sự khác biệt giữa hệ thống chưa bảo vệ và hệ thống đã tích hợp kiểm tra sống.",
            "Tiến độ chung đạt khoảng 95%, phần còn lại chủ yếu là chỉnh sửa câu chữ, kiểm tra hình ảnh/bảng số liệu và luyện trình bày.",
        ]),
        ("4. Khó khăn", [
            "Dữ liệu thử nghiệm chưa đủ lớn để khẳng định mô hình chống được mọi kiểu giả mạo nâng cao như deepfake chất lượng cao hoặc video replay phức tạp.",
            "Kết quả demo phụ thuộc vào điều kiện ánh sáng, chất lượng webcam và chất lượng ảnh giả mạo nên cần giải thích rõ phạm vi thực nghiệm.",
            "Cần trình bày rõ đây là môi trường lab phục vụ nghiên cứu và minh họa, chưa phải hệ thống thương mại hoàn chỉnh có đầy đủ bảo mật kênh truyền, kiểm soát thiết bị và giám sát vận hành.",
        ]),
        ("5. Kế hoạch tiếp theo", [
            "Hoàn thiện file báo cáo cuối, slide thuyết trình và video demo để nộp/trình bày.",
            "Kiểm tra lại chính tả, định dạng, hình ảnh, bảng số liệu và tên các thuật ngữ như FAR, FRR, APCER, Liveness Detection.",
            "Chuẩn bị trả lời các câu hỏi về điểm tấn công trong quy trình biometric, lý do đặt Liveness trước bước so khớp và hướng cải tiến mô hình.",
            "Đề xuất hướng phát triển tiếp theo: mở rộng dữ liệu, thử thêm video-based liveness, tăng bảo mật kênh truyền và kiểm thử với nhiều dạng spoofing hơn.",
        ]),
    ],
}


def set_paragraph_text(paragraph, text, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = bold


def replace_cell_content(cell, sections):
    base_paragraphs = list(cell.paragraphs)
    first = base_paragraphs[0]
    set_paragraph_text(first, sections[0][0], bold=True)
    for p in base_paragraphs[1:]:
        p._element.getparent().remove(p._element)

    for index, (section_title, bullets) in enumerate(sections):
        if index != 0:
            set_paragraph_text(cell.add_paragraph(), section_title, bold=True)
        for bullet in bullets:
            set_paragraph_text(cell.add_paragraph(), f"- {bullet}")


def create_report(week, sections):
    doc = Document(TEMPLATE)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("BÁO CÁO TIẾN ĐỘ"):
            set_paragraph_text(paragraph, f"BÁO CÁO TIẾN ĐỘ TUẦN {week}", bold=True)
            break

    replace_cell_content(doc.tables[0].cell(0, 0), sections)
    out = OUTPUT_DIR / f"BaoCaoTuan{week}.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    for week, sections in REPORTS.items():
        print(create_report(week, sections))
