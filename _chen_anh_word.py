# -*- coding: utf-8 -*-
"""
Chèn ảnh vào file BAOCAODOANCOSO_NGUYENPHUCTHINH.docx tại đúng vị trí
caption "Hình X.Y. ..." đã có sẵn, đồng thời chuẩn hoá định dạng:
  - Times New Roman, cỡ 13
  - Giãn dòng 1.5
  - Căn đều hai bên (Justify)
  - Lề chuẩn (2.54 cm tất cả)
  - Caption căn giữa, in đậm tiêu đề chương vẫn giữ
  - Bảo toàn bảng biểu, mục lục, số trang
"""
import os
import sys
import io
import re
import copy

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"C:/Users/ADMIN/biometric-security-lab"
ASSETS = os.path.join(ROOT, "_report_assets")
SRC = r"C:/Users/ADMIN/Downloads/BAOCAODOANCOSO_NGUYENPHUCTHINH.docx"
DST = r"C:/Users/ADMIN/Downloads/BAOCAODOANCOSO_NGUYENPHUCTHINH_DACHINH.docx"

# Map số hình ("2.1", "4.2b"...) -> file ảnh
FIG_MAP = {
    "2.1":  "fig_2_1_quy_trinh_bio.png",
    "2.2":  "fig_2_2_phan_loai_tan_cong.png",
    "2.3":  "fig_2_3_mobilenetv2.png",
    "2.4":  "fig_2_4_injection_pipeline.png",
    "2.5":  "fig_2_5_pa_vs_injection.png",
    "3.1":  "fig_3_1_kien_truc_tong_the.png",
    "3.2":  "fig_3_2_server_vulnerable.png",
    "3.3":  "fig_3_3_attack_replay_injection.png",
    "3.4":  "fig_3_4_injection_chi_tiet.png",
    "3.5":  "fig_3_5_tien_trinh_3_giai_doan.png",
    "4.1":  "fig_4_1_dataset.png",
    "4.2":  "fig_4_2_training_history.png",
    "4.2b": "fig_4_2b_confusion_matrix.png",
    "4.3":  "fig_4_3_server_secured.png",
    "4.4":  "fig_4_4_so_sanh_chi_so.png",
    "4.5":  "fig_4_5_lop_phong_ve.png",
    "4.6":  "fig_4_6_giao_dien_vulnerable.png",
    "4.7":  "fig_4_7_log_injection.png",
    "4.8":  "fig_4_8_chu_trinh.png",
    "5.1":  "fig_5_1_lo_trinh.png",
}

CAPTION_RE = re.compile(r"^Hình\s+(\d+\.\d+[a-z]?)\.\s", re.UNICODE)


def insert_paragraph_before(target_par):
    """Chèn một paragraph rỗng ngay TRƯỚC paragraph mục tiêu, trả về paragraph mới."""
    new_p = OxmlElement("w:p")
    target_par._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, target_par._parent)


def set_margins(doc, cm=2.54):
    for section in doc.sections:
        section.top_margin = Cm(cm)
        section.bottom_margin = Cm(cm)
        section.left_margin = Cm(cm)
        section.right_margin = Cm(cm)


def apply_font_to_run(run, name="Times New Roman", size_pt=13):
    run.font.name = name
    # Đảm bảo font tiếng Việt áp dụng cho cả East Asian / Complex Script
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def set_paragraph_format(par, justify=True, line_spacing=1.5):
    pf = par.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if justify and par.alignment is None:
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def normalize_paragraph(par, default_size=13, justify_normal=True):
    """Áp font/size cho mọi run, set giãn dòng & căn lề hợp lý theo style."""
    style_name = (par.style.name or "").lower()

    # Giữ alignment đặc biệt cho heading/caption/title
    is_heading = style_name.startswith("heading") or style_name.startswith("title")
    is_toc = style_name.startswith("toc")

    # Cỡ chữ heading lớn hơn
    if "heading 1" in style_name:
        size = 14
    elif "heading 2" in style_name:
        size = 13
    elif "heading 3" in style_name:
        size = 13
    else:
        size = default_size

    for run in par.runs:
        apply_font_to_run(run, "Times New Roman", size)

    pf = par.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5

    if not is_heading and not is_toc:
        if par.alignment is None or par.alignment == WD_ALIGN_PARAGRAPH.LEFT:
            # Caption hình sẽ được căn giữa ở bước sau
            if justify_normal:
                par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def find_caption_paragraphs(doc):
    """
    Trả về list (index, paragraph, fig_key) cho các paragraph caption "Hình X.Y. ..."
    nằm TRONG nội dung chương (sau heading "CHƯƠNG 1"), bỏ qua phần
    DANH MỤC HÌNH ẢNH.
    """
    paragraphs = doc.paragraphs
    # Tìm điểm bắt đầu nội dung: heading "CHƯƠNG 1"
    body_start = 0
    for i, p in enumerate(paragraphs):
        if p.style.name.startswith("Heading") and "CHƯƠNG 1" in p.text.upper():
            body_start = i
            break

    found = []
    for i in range(body_start, len(paragraphs)):
        par = paragraphs[i]
        text = par.text.strip()
        m = CAPTION_RE.match(text)
        if m:
            key = m.group(1)
            if key in FIG_MAP:
                found.append((i, par, key))
    return found


def insert_image_before_caption(par, image_path, max_width_inches=6.0):
    """
    Chèn 1 paragraph mới chứa ảnh ngay TRƯỚC caption.
    Ảnh căn giữa, rộng tối đa max_width_inches inches.
    """
    img_par = insert_paragraph_before(par)
    img_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_par.add_run()
    run.add_picture(image_path, width=Inches(max_width_inches))
    # Khoảng cách trên/dưới nhỏ
    pf = img_par.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return img_par


def style_caption(par):
    """Caption: căn giữa, in nghiêng, font Times New Roman 12, đậm 'Hình X.Y.'"""
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = par.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    for run in par.runs:
        apply_font_to_run(run, "Times New Roman", 12)
        run.italic = True


def main():
    print(f"Mở: {SRC}")
    doc = Document(SRC)

    # 1) Lề chuẩn
    set_margins(doc, cm=2.54)

    # 2) Tìm caption rồi chèn ảnh (duyệt từ cuối lên đầu để index không lệch)
    captions = find_caption_paragraphs(doc)
    print(f"Tìm thấy {len(captions)} caption hình ảnh trong phần nội dung.")

    inserted = 0
    missing = []
    for idx, par, key in reversed(captions):
        fname = FIG_MAP.get(key)
        if not fname:
            continue
        fpath = os.path.join(ASSETS, fname)
        if not os.path.exists(fpath):
            missing.append((key, fpath))
            continue
        insert_image_before_caption(par, fpath, max_width_inches=6.0)
        style_caption(par)
        inserted += 1

    print(f"Đã chèn {inserted} ảnh.")
    if missing:
        print("Thiếu file ảnh:")
        for k, p in missing:
            print(f"  - Hình {k} -> {p}")

    # 3) Chuẩn hoá định dạng toàn bộ paragraph (sau khi chèn xong để bao luôn các paragraph mới)
    for par in doc.paragraphs:
        # Bỏ qua paragraph chỉ chứa ảnh (không có run text) — đã set ở insert
        text = par.text.strip()
        if not text:
            # vẫn set line spacing đẹp hơn cho khoảng trắng
            pf = par.paragraph_format
            if pf.line_spacing_rule is None:
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            continue
        # Caption hình ảnh: giữ kiểu căn giữa
        if CAPTION_RE.match(text):
            style_caption(par)
            continue
        normalize_paragraph(par)

    # 4) Định dạng cell trong bảng
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    for run in par.runs:
                        apply_font_to_run(run, "Times New Roman", 12)
                    pf = par.paragraph_format
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = 1.3

    # 5) Cập nhật style 'Normal' để mặc định khớp định dạng
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Times New Roman"
    style_normal.font.size = Pt(13)
    rPr = style_normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), "Times New Roman")

    doc.save(DST)
    print(f"Đã lưu: {DST}")


if __name__ == "__main__":
    main()
