# -*- coding: utf-8 -*-
"""Build BAOCAODOANCOSO_NGUYENPHUCTHINH_FINAL.docx theo yeu cau chi tiet."""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASSETS = r"C:\Users\ADMIN\biometric-security-lab\_report_assets"
OUT_PATH = r"C:\Users\ADMIN\Downloads\BAOCAODOANCOSO_NGUYENPHUCTHINH_FINAL.docx"

# Quan ly so trang du kien cho Danh muc hinh / bang
FIGURES = []  # list of dict: {"num":"1.1","title":"...","page":12}
TABLES = []   # list of dict: {"num":"2.1","title":"...","page":18}


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        v = kwargs.get(edge, {'sz': '6', 'val': 'single', 'color': '000000'})
        e = OxmlElement(f'w:{edge}')
        for k, val in v.items():
            e.set(qn(f'w:{k}'), val)
        tcBorders.append(e)
    tcPr.append(tcBorders)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(6)  # WD_BREAK.PAGE = 6


def set_run_default(run, size=13, bold=False, italic=False, font='Times New Roman'):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, size=13, bold=False, italic=False,
             align='justify', indent_first=0.0, space_after=6, level=None):
    p = doc.add_paragraph()
    if level is not None:
        # heading style
        p.style = doc.styles['Normal']
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    run = p.add_run(text)
    set_run_default(run, size=size, bold=bold, italic=italic)
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text.upper())
    set_run_default(run, size=14, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_default(run, size=13, bold=True)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_default(run, size=13, bold=True, italic=True)
    return p


def add_bullet(doc, text, *, size=13):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run('')
    run.text = text
    set_run_default(run, size=size)
    return p


def add_image(doc, filename, width_cm=14, fig_num=None, caption=None, page_hint=None):
    path = os.path.join(ASSETS, filename) if not os.path.isabs(filename) else filename
    if not os.path.exists(path):
        # bo qua hinh thieu
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption is not None and fig_num is not None:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        r = cap.add_run(f"Hình {fig_num}. {caption}")
        set_run_default(r, size=12, italic=True, bold=True)
        FIGURES.append({"num": fig_num, "title": caption, "page": page_hint or "—"})


def add_table_caption(doc, num, title, page_hint=None):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(f"Bảng {num}. {title}")
    set_run_default(r, size=12, italic=True, bold=True)
    TABLES.append({"num": num, "title": title, "page": page_hint or "—"})


def add_basic_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        para = c.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h)
        set_run_default(r, size=12, bold=True)
        shade_cell(c, 'D9E1F2')
        set_cell_border(c)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = table.rows[ri].cells[ci]
            c.text = ""
            para = c.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.2
            r = para.add_run(val)
            set_run_default(r, size=12)
            set_cell_border(c)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def setup_document_defaults(doc):
    # margin
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
    # default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), 'Times New Roman')


def build_cover(doc):
    add_para(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", size=13, bold=True, align='center', space_after=2)
    add_para(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HCM (HUTECH)",
             size=13, bold=True, align='center', space_after=2)
    add_para(doc, "KHOA CÔNG NGHỆ THÔNG TIN", size=13, bold=True, align='center', space_after=20)
    add_para(doc, "ĐỒ ÁN CƠ SỞ", size=22, bold=True, align='center', space_after=10)
    add_para(doc,
             "NGHIÊN CỨU TẤN CÔNG VÀO HỆ THỐNG\nSINH TRẮC HỌC (BIOMETRIC SECURITY)",
             size=18, bold=True, align='center', space_after=20)
    add_para(doc, "Ngành: AN TOÀN THÔNG TIN", size=14, bold=True, align='center', space_after=14)
    add_para(doc, "Giảng viên hướng dẫn: Cô Đặng Thị Thạch Thảo",
             size=13, align='center', space_after=4)
    add_para(doc, "Sinh viên thực hiện: Nguyễn Phúc Thịnh",
             size=13, align='center', space_after=4)
    add_para(doc, "MSSV: 2387700066    Lớp: 23DATA1",
             size=13, align='center', space_after=20)
    add_para(doc, "TP. Hồ Chí Minh, năm 2026",
             size=13, italic=True, align='center')
    add_page_break(doc)


def build_front_matter(doc):
    # LOI MO DAU
    add_heading1(doc, "LỜI MỞ ĐẦU")
    add_para(doc,
             "Trong hơn một thập kỷ trở lại đây, xác thực sinh trắc học (biometric authentication) "
             "đã trở thành phương thức xác thực chủ đạo, được tích hợp rộng rãi trên điện thoại thông minh, "
             "hệ thống ngân hàng điện tử, kiểm soát ra vào và các dịch vụ định danh công dân điện tử. "
             "So với mật khẩu truyền thống, sinh trắc học mang lại trải nghiệm thuận tiện hơn vì người dùng "
             "không cần ghi nhớ thông tin và tốc độ xác thực nhanh chỉ trong vài trăm mili-giây [1][12].",
             indent_first=1.0)
    add_para(doc,
             "Tuy nhiên, đi cùng với sự phổ biến đó là các nguy cơ bảo mật ngày càng nghiêm trọng. "
             "Các hệ thống nhận diện khuôn mặt nếu không được trang bị cơ chế kiểm tra sống "
             "(Liveness Detection / Presentation Attack Detection) rất dễ bị qua mặt bằng ảnh in, "
             "ảnh hiển thị trên điện thoại, video phát lại hoặc các kỹ thuật injection trực tiếp ở tầng API. "
             "Khác với mật khẩu, dữ liệu sinh trắc học một khi bị lộ thì không thể thay đổi, "
             "vì vậy hậu quả của một cuộc tấn công thành công là rất lâu dài [1][9][29].",
             indent_first=1.0)
    add_para(doc,
             "Đề tài “Nghiên cứu tấn công vào hệ thống sinh trắc học” được thực hiện nhằm khảo sát "
             "các bề mặt tấn công của một hệ thống xác thực khuôn mặt thực tế, xây dựng môi trường "
             "thí nghiệm gồm server có lỗ hổng và server đã được gia cố, từ đó đánh giá định lượng "
             "hiệu quả của giải pháp Liveness Detection dựa trên mạng MobileNetV2 [3]. Báo cáo trình bày "
             "toàn bộ quá trình từ cơ sở lý thuyết, triển khai, tấn công, phòng vệ cho đến đánh giá kết quả, "
             "hướng đến mục tiêu cảnh báo rủi ro và đề xuất biện pháp giảm thiểu cho các hệ thống "
             "sinh trắc học trong thực tế.",
             indent_first=1.0)
    add_page_break(doc)

    # LOI CAM DOAN
    add_heading1(doc, "LỜI CAM ĐOAN")
    add_para(doc,
             "Em xin cam đoan đồ án “Nghiên cứu tấn công vào hệ thống sinh trắc học” là công trình "
             "nghiên cứu của riêng em, được thực hiện dưới sự hướng dẫn của giảng viên Đặng Thị Thạch Thảo. "
             "Các số liệu, kết quả thực nghiệm trong báo cáo là trung thực, được đo đạc trên môi trường "
             "lab do em tự xây dựng và chưa từng được công bố trong bất kỳ công trình nào khác.",
             indent_first=1.0)
    add_para(doc,
             "Các nguồn tài liệu tham khảo đều được trích dẫn rõ ràng theo định dạng [n]. Toàn bộ "
             "thí nghiệm tấn công chỉ được thực hiện trên hệ thống do em tự triển khai phục vụ mục đích "
             "nghiên cứu, không nhằm vào bất kỳ hệ thống thương mại của tổ chức hoặc cá nhân khác.",
             indent_first=1.0)
    add_para(doc, "TP. HCM, ngày … tháng … năm 2026", align='right', italic=True, space_after=2)
    add_para(doc, "Sinh viên thực hiện", align='right', italic=True, space_after=20)
    add_para(doc, "Nguyễn Phúc Thịnh", align='right', bold=True)
    add_page_break(doc)

    # LOI CAM ON - can trai
    add_heading1(doc, "LỜI CẢM ƠN")
    add_para(doc,
             "Trước tiên, em xin gửi lời cảm ơn chân thành đến quý thầy cô Khoa Công nghệ Thông tin – "
             "Trường Đại học Công nghệ TP. HCM (HUTECH) đã trang bị cho em những kiến thức nền tảng "
             "quan trọng về an toàn thông tin, học máy và phát triển ứng dụng trong suốt thời gian học tập "
             "tại trường.",
             align='left', indent_first=1.0)
    add_para(doc,
             "Em xin gửi lời cảm ơn sâu sắc nhất đến giảng viên hướng dẫn Cô Đặng Thị Thạch Thảo "
             "đã tận tình góp ý, định hướng và tạo điều kiện để em có thể hoàn thành đề tài. Những phản hồi "
             "kịp thời của Cô trong từng tuần báo cáo đã giúp em điều chỉnh hướng nghiên cứu phù hợp và "
             "bám sát mục tiêu đặt ra ban đầu.",
             align='left', indent_first=1.0)
    add_para(doc,
             "Em cũng xin cảm ơn gia đình, bạn bè đã luôn động viên, hỗ trợ em trong quá trình nghiên cứu "
             "và thử nghiệm các kịch bản tấn công, phòng vệ trên hệ thống sinh trắc học.",
             align='left', indent_first=1.0)
    add_para(doc,
             "Do thời gian và kiến thức còn hạn chế, đồ án không tránh khỏi những thiếu sót. "
             "Em rất mong nhận được những góp ý quý báu từ quý thầy cô để hoàn thiện hơn trong các "
             "nghiên cứu tiếp theo. Em xin chân thành cảm ơn!",
             align='left', indent_first=1.0)
    add_page_break(doc)

    # NHAN XET GVHD
    add_heading1(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN")
    for _ in range(10):
        add_para(doc, "………………………………………………………………………………………………",
                 align='left', space_after=4)
    add_para(doc, "TP. HCM, ngày … tháng … năm 2026", align='right', italic=True, space_after=2)
    add_para(doc, "Giảng viên hướng dẫn", align='right', italic=True, bold=True)
    add_page_break(doc)


def build_toc(doc):
    add_heading1(doc, "MỤC LỤC")
    toc_entries = [
        ("LỜI MỞ ĐẦU", "i", 0),
        ("LỜI CAM ĐOAN", "ii", 0),
        ("LỜI CẢM ƠN", "iii", 0),
        ("NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", "iv", 0),
        ("MỤC LỤC", "v", 0),
        ("DANH MỤC CHỮ VIẾT TẮT", "vii", 0),
        ("DANH MỤC HÌNH ẢNH", "viii", 0),
        ("DANH MỤC BẢNG BIỂU", "ix", 0),
        ("CHƯƠNG 1: MỞ ĐẦU", "1", 0),
        ("1.1. Lý do chọn đề tài", "1", 1),
        ("1.2. Mục tiêu nghiên cứu", "5", 1),
        ("1.3. Đối tượng và phạm vi nghiên cứu", "5", 1),
        ("1.3.1. Đối tượng nghiên cứu", "5", 2),
        ("1.3.2. Phạm vi nghiên cứu", "6", 2),
        ("1.4. Ý nghĩa khoa học và thực tiễn", "6", 1),
        ("1.5. Dự kiến kết quả đạt được", "7", 1),
        ("1.6. Bố cục báo cáo", "8", 1),
        ("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT", "9", 0),
        ("2.1. Tổng quan về nhận diện khuôn mặt", "9", 1),
        ("2.1.1. Khái niệm Face Detection", "9", 2),
        ("2.1.2. Khái niệm Face Recognition", "10", 2),
        ("2.1.3. Khái niệm Face Liveness Detection", "11", 2),
        ("2.1.4. Sinh trắc học khuôn mặt", "12", 2),
        ("2.2. Vai trò của bảo mật trong hệ thống nhận diện khuôn mặt", "13", 1),
        ("2.3. Các dạng dữ liệu đầu vào", "14", 1),
        ("2.4. Quy trình chung của hệ thống nhận diện khuôn mặt", "15", 1),
        ("2.5. Các mô hình kiến trúc tiêu biểu", "17", 1),
        ("2.5.1. CNN – Convolutional Neural Network", "17", 2),
        ("2.5.2. MobileNetV2", "18", 2),
        ("2.5.3. ResNet", "19", 2),
        ("2.5.4. VGGFace", "20", 2),
        ("2.5.5. FaceNet", "21", 2),
        ("2.5.6. MTCNN và Haar Cascade", "22", 2),
        ("2.6. Các loại tấn công vào hệ thống nhận diện khuôn mặt", "24", 1),
        ("2.6.1. Tấn công bằng ảnh in", "24", 2),
        ("2.6.2. Tấn công bằng ảnh trên điện thoại / màn hình", "25", 2),
        ("2.6.3. Tấn công bằng video phát lại (Replay)", "25", 2),
        ("2.6.4. Tấn công bằng mặt nạ 3D", "26", 2),
        ("2.6.5. Tấn công bằng Deepfake", "26", 2),
        ("2.6.6. Tấn công bằng camera ảo / video giả lập", "27", 2),
        ("2.6.7. Tấn công Injection ở tầng API", "28", 2),
        ("2.7. Quá trình tấn công chi tiết", "29", 1),
        ("2.8. Tính chất bảo mật trong hệ thống nhận diện khuôn mặt", "30", 1),
        ("2.9. Liveness Detection và tiêu chuẩn ISO/IEC 30107", "32", 1),
        ("2.10. Các chỉ số đánh giá", "33", 1),
        ("2.11. OWASP Biometric Risks", "34", 1),
        ("CHƯƠNG 3: PHÂN TÍCH MÔ HÌNH / HỆ THỐNG", "36", 0),
        ("3.1. Mô hình kiến trúc tổng thể", "36", 1),
        ("3.2. Phiên bản server có lỗ hổng (vulnerable)", "37", 1),
        ("3.2.1. Mã nguồn cốt lõi của /authenticate", "37", 2),
        ("3.2.2. Bề mặt tấn công", "38", 2),
        ("3.3. Phiên bản server đã gia cố (secured)", "39", 1),
        ("3.4. So sánh hai phiên bản server", "40", 1),
        ("3.5. Tiền xử lý ảnh đầu vào", "41", 1),
        ("3.6. Mô hình đe dọa (Threat Model)", "42", 1),
        ("CHƯƠNG 4: MÔ PHỎNG, ĐÁNH GIÁ VÀ KẾT QUẢ", "44", 0),
        ("4.1. Môi trường thí nghiệm", "44", 1),
        ("4.2. Kịch bản tấn công trên server vulnerable", "45", 1),
        ("4.2.1. Spoofing bằng ảnh in", "45", 2),
        ("4.2.2. Spoofing bằng ảnh trên điện thoại", "46", 2),
        ("4.2.3. Tấn công Replay qua API", "46", 2),
        ("4.2.4. Tấn công Injection", "47", 2),
        ("4.3. Xây dựng giải pháp Liveness Detection", "48", 1),
        ("4.3.1. Bộ dữ liệu", "48", 2),
        ("4.3.2. Cấu hình huấn luyện", "49", 2),
        ("4.3.3. Tích hợp vào server secured", "49", 2),
        ("4.4. Kết quả đánh giá định lượng", "50", 1),
        ("4.5. Phân tích kết quả", "51", 1),
        ("4.6. Mô phỏng quy trình phòng chống tấn công", "52", 1),
        ("4.7. Hạn chế của giải pháp", "53", 1),
        ("KẾT LUẬN", "54", 0),
        ("TÀI LIỆU THAM KHẢO", "56", 0),
    ]
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    for text, page, lvl in toc_entries:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(2)
        pf.line_spacing = 1.3
        if lvl == 0:
            pf.left_indent = Cm(0)
            r = p.add_run(text)
            set_run_default(r, size=12, bold=True)
        elif lvl == 1:
            pf.left_indent = Cm(0.6)
            r = p.add_run(text)
            set_run_default(r, size=12)
        else:
            pf.left_indent = Cm(1.2)
            r = p.add_run(text)
            set_run_default(r, size=12)
        tab_run = p.add_run("\t")
        set_run_default(tab_run, size=12)
        page_run = p.add_run(str(page))
        set_run_default(page_run, size=12, bold=(lvl == 0))
        pf.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    add_page_break(doc)


def build_abbreviations(doc):
    add_heading1(doc, "DANH MỤC CHỮ VIẾT TẮT")
    headers = ["STT", "Viết tắt", "Nghĩa đầy đủ"]
    rows = [
        ("1", "FAR", "False Acceptance Rate – Tỷ lệ chấp nhận sai"),
        ("2", "FRR", "False Rejection Rate – Tỷ lệ từ chối sai"),
        ("3", "EER", "Equal Error Rate – Tỷ lệ lỗi cân bằng"),
        ("4", "APCER", "Attack Presentation Classification Error Rate"),
        ("5", "BPCER", "Bona-fide Presentation Classification Error Rate"),
        ("6", "PAD", "Presentation Attack Detection – Phát hiện tấn công trình diện"),
        ("7", "PAI", "Presentation Attack Instrument – Công cụ tấn công trình diện"),
        ("8", "CNN", "Convolutional Neural Network – Mạng nơ-ron tích chập"),
        ("9", "FR", "Face Recognition – Nhận diện khuôn mặt"),
        ("10", "FD", "Face Detection – Phát hiện khuôn mặt"),
        ("11", "MTCNN", "Multi-task Cascaded CNN – Phát hiện khuôn mặt đa tác vụ"),
        ("12", "MobileNetV2", "Kiến trúc CNN nhẹ dùng cho thiết bị di động"),
        ("13", "ResNet", "Residual Network – Mạng học sâu có kết nối tắt"),
        ("14", "VGGFace", "Mô hình nhận diện khuôn mặt của VGG, ĐH Oxford"),
        ("15", "FaceNet", "Mô hình nhận diện khuôn mặt của Google, embedding 128-d"),
        ("16", "GAN", "Generative Adversarial Network – Mạng đối kháng sinh"),
        ("17", "ISO/IEC 30107", "Tiêu chuẩn quốc tế về phát hiện tấn công sinh trắc học"),
        ("18", "API", "Application Programming Interface"),
        ("19", "REST", "Representational State Transfer"),
        ("20", "JSON", "JavaScript Object Notation"),
        ("21", "HTTP/HTTPS", "Giao thức truyền tải siêu văn bản (có/không mã hóa)"),
        ("22", "TLS", "Transport Layer Security – Mã hóa kênh truyền"),
        ("23", "OWASP", "Open Worldwide Application Security Project"),
        ("24", "CIA", "Confidentiality – Integrity – Availability"),
        ("25", "ROI", "Region of Interest – Vùng quan tâm trên ảnh"),
        ("26", "RGB", "Red – Green – Blue"),
        ("27", "IR", "Infrared – Hồng ngoại"),
        ("28", "GPU", "Graphics Processing Unit"),
        ("29", "MFA", "Multi-Factor Authentication – Xác thực đa yếu tố"),
    ]
    add_basic_table(doc, headers, rows, widths=[1.5, 3.0, 11])
    add_page_break(doc)


def build_list_of_figures(doc):
    add_heading1(doc, "DANH MỤC HÌNH ẢNH")
    headers = ["Tên hình", "Trang"]
    rows = [(f"Hình {f['num']}. {f['title']}", str(f['page'])) for f in FIGURES]
    if not rows:
        rows = [("(Sẽ được tự động cập nhật khi build báo cáo)", "—")]
    add_basic_table(doc, headers, rows, widths=[14, 2])
    add_page_break(doc)


def build_list_of_tables(doc):
    add_heading1(doc, "DANH MỤC BẢNG BIỂU")
    headers = ["Tên bảng", "Trang"]
    rows = [(f"Bảng {t['num']}. {t['title']}", str(t['page'])) for t in TABLES]
    if not rows:
        rows = [("(Sẽ được tự động cập nhật khi build báo cáo)", "—")]
    add_basic_table(doc, headers, rows, widths=[14, 2])
    add_page_break(doc)


# ============================================================
# CHUONG 1
# ============================================================
def build_chapter1(doc):
    add_heading1(doc, "CHƯƠNG 1: MỞ ĐẦU")

    add_heading2(doc, "1.1. Lý do chọn đề tài")
    add_para(doc,
        "Trong bối cảnh chuyển đổi số diễn ra mạnh mẽ trên toàn cầu, xác thực danh tính người dùng "
        "đã trở thành một trong những thách thức cốt lõi của an toàn thông tin. Các phương thức xác thực "
        "truyền thống như mật khẩu hay mã PIN ngày càng bộc lộ nhiều hạn chế: dễ bị đánh cắp, quên mất "
        "hoặc bị tấn công brute-force. Trước thực trạng đó, xác thực sinh trắc học – đặc biệt là nhận diện "
        "khuôn mặt – đã nổi lên như một giải pháp thay thế ưu việt nhờ tính tiện lợi, tốc độ và khả năng "
        "định danh dựa trên đặc trưng vật lý vốn có của mỗi cá nhân [1][26].",
        indent_first=1.0)
    add_para(doc,
        "Theo thống kê của Statista (2023), hơn 1,4 tỷ thiết bị di động trên thế giới đã tích hợp "
        "tính năng nhận diện khuôn mặt. Tại Việt Nam, từ năm 2024, Ngân hàng Nhà nước yêu cầu các tổ chức "
        "tín dụng áp dụng xác thực sinh trắc học cho các giao dịch chuyển khoản trên 10 triệu đồng, "
        "đánh dấu bước chuyển quan trọng trong việc ứng dụng công nghệ này vào hệ thống tài chính quốc gia. "
        "Bên cạnh đó, hệ thống kiểm soát ra vào tại các tòa nhà, sân bay, trường học và cơ quan nhà nước "
        "cũng đang dần thay thế thẻ từ bằng nhận diện khuôn mặt [12][29].",
        indent_first=1.0)
    add_para(doc,
        "Tuy nhiên, song song với sự phổ biến đó là những nguy cơ bảo mật ngày càng tinh vi và nguy hiểm. "
        "Nhiều nghiên cứu độc lập đã chứng minh rằng phần lớn hệ thống nhận diện khuôn mặt thế hệ đầu "
        "chỉ thực hiện so khớp đặc trưng (feature matching) mà hoàn toàn bỏ qua bước kiểm tra xem khuôn "
        "mặt được trình diện có phải từ người sống thật hay không. Điều này tạo ra một lỗ hổng nghiêm trọng: "
        "kẻ tấn công chỉ cần một tấm ảnh chân dung của nạn nhân – dễ dàng thu thập từ mạng xã hội – "
        "là có thể qua mặt hệ thống [6][21][22].",
        indent_first=1.0)

    add_image(doc, "fig_2_2_phan_loai_tan_cong.png", width_cm=14,
              fig_num="1.1",
              caption="Phân loại các dạng tấn công vào hệ thống nhận diện khuôn mặt",
              page_hint="2")
    add_para(doc,
        "Hình 1.1 minh họa các dạng tấn công phổ biến nhất vào hệ thống nhận diện khuôn mặt, "
        "được phân loại theo tiêu chuẩn ISO/IEC 30107 [1]. Từ trái sang phải: tấn công bằng ảnh in "
        "(Print Attack), ảnh hiển thị trên màn hình điện thoại (Screen Replay), video phát lại "
        "(Video Replay), mặt nạ 3D (3D Mask) và tấn công Deepfake sử dụng mạng GAN. Mỗi dạng tấn công "
        "đặt ra yêu cầu phòng vệ khác nhau, đòi hỏi hệ thống phải có cơ chế Liveness Detection đủ mạnh.",
        indent_first=1.0)
    add_para(doc,
        "Đặc biệt đáng lo ngại là sự xuất hiện của công nghệ Deepfake – sử dụng mạng đối kháng sinh "
        "(GAN) để tổng hợp video khuôn mặt giả với biểu cảm sống động, gần như không thể phân biệt "
        "bằng mắt thường [16][17][18]. Các công cụ tạo Deepfake ngày càng dễ tiếp cận, thậm chí có "
        "ứng dụng miễn phí trên điện thoại, khiến ngưỡng kỹ thuật để thực hiện tấn công ngày càng thấp "
        "trong khi mức độ nguy hiểm ngày càng cao. Bên cạnh đó, tấn công Injection – gửi trực tiếp ảnh "
        "hoặc dữ liệu đặc trưng vào API xác thực mà không cần thiết bị vật lý – là dạng tấn công "
        "hoàn toàn vô hình với các cơ chế bảo vệ truyền thống [9][11].",
        indent_first=1.0)

    add_image(doc, "fig_2_4_injection_pipeline.png", width_cm=14,
              fig_num="1.2",
              caption="Các điểm tấn công Injection trong pipeline xác thực sinh trắc học",
              page_hint="3")
    add_para(doc,
        "Hình 1.2 thể hiện các vị trí mà kẻ tấn công có thể can thiệp vào pipeline xác thực sinh trắc học "
        "theo mô hình ISO/IEC 30107 [1]. Điểm tấn công (1) là trước cảm biến – kẻ tấn công đưa PAI "
        "(Presentation Attack Instrument) vào trực tiếp trước camera. Điểm (2) là giữa cảm biến và "
        "module xử lý – dữ liệu bị chặn và thay thế trên đường truyền. Điểm (3) là tại tầng API – "
        "kẻ tấn công gửi ảnh hoặc vector đặc trưng giả trực tiếp vào endpoint xác thực, "
        "hoàn toàn bypass cảm biến vật lý.",
        indent_first=1.0)
    add_para(doc,
        "Vấn đề đặt ra là cần xây dựng và nghiên cứu các giải pháp giúp hệ thống phân biệt được "
        "khuôn mặt thật và khuôn mặt giả mạo, đồng thời đánh giá định lượng hiệu quả của các giải pháp "
        "đó trên môi trường thực nghiệm có kiểm soát. Đây chính là động lực để đề tài "
        "“Nghiên cứu tấn công vào hệ thống sinh trắc học” được thực hiện.",
        indent_first=1.0)
    add_para(doc,
        "Về phương pháp tiếp cận, đề tài kết hợp ba hướng nghiên cứu: (1) nghiên cứu lý thuyết – "
        "tổng hợp tài liệu, tiêu chuẩn ISO/IEC 30107 và các công bố khoa học về Liveness Detection; "
        "(2) thực nghiệm – xây dựng hệ thống, mô phỏng tấn công và thu thập số liệu; "
        "(3) so sánh đánh giá – đối chiếu kết quả của hệ thống có và không có Liveness Detection "
        "trên cùng bộ dữ liệu kiểm thử [2][9].",
        indent_first=1.0)

    add_heading2(doc, "1.2. Mục tiêu nghiên cứu")
    add_bullet(doc, "Hệ thống hóa cơ sở lý thuyết về xác thực sinh trắc học, các bề mặt tấn công và tiêu chuẩn ISO/IEC 30107 [1][2].")
    add_bullet(doc, "Xây dựng môi trường thí nghiệm gồm hai phiên bản server: bản dễ bị tổn thương (vulnerable) và bản đã gia cố (secured) bằng Liveness Detection.")
    add_bullet(doc, "Triển khai và phân tích các kịch bản tấn công thực tế: spoofing bằng ảnh in, ảnh trên điện thoại, replay, injection trực tiếp qua API.")
    add_bullet(doc, "Huấn luyện mô hình phân loại Real/Fake dựa trên MobileNetV2 [3] với dataset tự thu thập gồm 200 ảnh thật và 200 ảnh giả.")
    add_bullet(doc, "Đánh giá định lượng bằng các chỉ số FAR, FRR, APCER, BPCER trước và sau khi tích hợp giải pháp phòng vệ [1][2].")

    add_heading2(doc, "1.3. Đối tượng và phạm vi nghiên cứu")
    add_heading3(doc, "1.3.1. Đối tượng nghiên cứu")
    add_para(doc,
        "Đối tượng nghiên cứu là hệ thống xác thực bằng khuôn mặt (face authentication) được triển khai "
        "dưới dạng ứng dụng web/API, sử dụng thư viện face_recognition (dlib) [4][5] cho nhận diện và "
        "mạng MobileNetV2 [3] cho phát hiện tấn công trình diện.",
        indent_first=1.0)
    add_heading3(doc, "1.3.2. Phạm vi nghiên cứu")
    add_bullet(doc, "Trong phạm vi lab: chỉ tấn công và phòng vệ trên hệ thống do nhóm tự xây dựng, không nhằm vào bất kỳ hệ thống thương mại nào.")
    add_bullet(doc, "Tập trung vào sinh trắc khuôn mặt 2D từ camera RGB; không khảo sát IR, 3D, vân tay hay mống mắt.")
    add_bullet(doc, "Các kịch bản tấn công giới hạn ở: ảnh in giấy, ảnh hiển thị trên điện thoại, ảnh chụp lại màn hình và injection ảnh tĩnh qua API.")

    add_heading2(doc, "1.4. Ý nghĩa khoa học và thực tiễn")
    add_para(doc,
        "Về mặt khoa học, đề tài hệ thống lại các kịch bản tấn công sinh trắc theo đúng vị trí trong "
        "quy trình xác thực và đề xuất vị trí phù hợp để đặt cơ chế phòng vệ. Kết quả thực nghiệm "
        "cung cấp bộ số liệu định lượng (FAR, FRR, APCER, BPCER) có thể tái sử dụng cho các nghiên cứu "
        "tiếp theo trong lĩnh vực PAD [9][21].",
        indent_first=1.0)
    add_para(doc,
        "Về mặt thực tiễn, đề tài cung cấp một bộ tài liệu mở gồm mã nguồn demo, dataset và hướng dẫn "
        "triển khai có thể tái sử dụng cho mục đích đào tạo và đánh giá an toàn cho các hệ thống biometric "
        "đang vận hành. Đây cũng là tài liệu tham khảo hữu ích cho các tổ chức đang triển khai hoặc "
        "nâng cấp hệ thống xác thực khuôn mặt [11][30].",
        indent_first=1.0)

    add_heading2(doc, "1.5. Dự kiến kết quả đạt được")
    add_para(doc, "Sau khi hoàn thành đề tài, các kết quả dự kiến đạt được bao gồm:", indent_first=1.0)
    add_bullet(doc, "Hiểu được nguyên lý hoạt động của hệ thống nhận diện khuôn mặt, bao gồm các bước từ phát hiện khuôn mặt, trích xuất đặc trưng đến so khớp và ra quyết định xác thực.")
    add_bullet(doc, "Nắm vững khái niệm và cơ chế của Liveness Detection (Presentation Attack Detection) theo tiêu chuẩn ISO/IEC 30107 [1][2].")
    add_bullet(doc, "Xây dựng được quy trình xử lý ảnh/khuôn mặt hoàn chỉnh: từ thu nhận dữ liệu, tiền xử lý, phát hiện khuôn mặt đến phân loại Real/Fake.")
    add_bullet(doc, "Phân tích được các kiểu tấn công thường gặp: Print Attack, Screen Replay, Video Replay, 3D Mask, Deepfake và Injection Attack.")
    add_bullet(doc, "Đánh giá được ưu điểm và nhược điểm của MobileNetV2 [3] trong bài toán Liveness Detection, so sánh với ResNet [13], VGGFace [15], FaceNet [14].")
    add_bullet(doc, "Đề xuất hướng cải thiện bảo mật: kết hợp MFA, xác thực thiết bị, giám sát vận hành và mở rộng sang Liveness Detection dựa trên video.")
    add_bullet(doc, "Cung cấp bộ số liệu thực nghiệm (FAR ≈ 2%, FRR ≈ 5%, APCER ≈ 1.5%, BPCER ≈ 4.8%) làm cơ sở so sánh cho các nghiên cứu tiếp theo.")

    add_heading2(doc, "1.6. Bố cục báo cáo")
    add_para(doc, "Báo cáo được tổ chức thành 4 chương chính:", indent_first=1.0)
    add_bullet(doc, "Chương 1 – Mở đầu: lý do chọn đề tài, mục tiêu, đối tượng, phạm vi, ý nghĩa và dự kiến kết quả.")
    add_bullet(doc, "Chương 2 – Cơ sở lý thuyết: các khái niệm về nhận diện khuôn mặt, sinh trắc học, mô hình kiến trúc, các dạng tấn công và tính chất bảo mật.")
    add_bullet(doc, "Chương 3 – Phân tích mô hình/hệ thống: kiến trúc hệ thống thí nghiệm, hai phiên bản server và mô hình đe dọa.")
    add_bullet(doc, "Chương 4 – Mô phỏng, đánh giá và kết quả: các kịch bản tấn công, giải pháp Liveness Detection, kết quả định lượng và phân tích.")
    add_page_break(doc)
