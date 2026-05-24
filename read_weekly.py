# -*- coding: utf-8 -*-
import sys
from pathlib import Path
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

DOWNLOADS = Path(r"C:\Users\ADMIN\Downloads")

for week in [3, 4, 5, 6]:
    f = DOWNLOADS / f"BaoCaoTuan{week}.docx"
    if not f.exists():
        continue
    print("=" * 80)
    print(f"=== WEEK {week} ===")
    print("=" * 80)
    doc = Document(f)
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text.strip())
    for ti, t in enumerate(doc.tables):
        print(f"--- Table {ti} ---")
        for row in t.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    print(txt)
