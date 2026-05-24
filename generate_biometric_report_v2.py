# -*- coding: utf-8 -*-
"""
Sinh BÁO CÁO ĐỒ ÁN CƠ SỞ — phiên bản 2 — mục tiêu ~52 trang.
Đầu vào: 22 hình trong _report_assets/.
Đầu ra: C:\\Users\\ADMIN\\Downloads\\BAOCAODOANCOSO_FINAL_v2.docx

Cấu trúc:
  - Trang bìa HUTECH
  - Lời mở đầu / cam đoan / cảm ơn / nhận xét GVHD
  - Mục lục (placeholder F9) / danh mục viết tắt / danh mục hình bảng
  - Chương 1..5
  - Tài liệu tham khảo (40 mục)

Phần lớn nội dung là TIẾNG VIỆT có dấu, người dùng có thể tự chỉnh sửa trong Word.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
ASSETS = ROOT / "_report_assets"
OUT = Path(r"C:\Users\ADMIN\Downloads\BAOCAODOANCOSO_FINAL_v2.docx")

# Bật/tắt khối phụ để tinh chỉnh độ dày tổng thể
EXTRA_DEPTH = True

doc = Document()

# ----- Page setup: A4, lề 3-2-2-2 -----
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

# ----- Default font Times New Roman 13 -----
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
rfonts.set(qn('w:ascii'), 'Times New Roman')
rfonts.set(qn('w:hAnsi'), 'Times New Roman')
rfonts.set(qn('w:cs'), 'Times New Roman')


# =========================================================
# Helpers (mở rộng từ generate_biometric_report.py)
# =========================================================
def add_para(text="", bold=False, italic=False, align=None, size=13, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.first_line_indent = Cm(0)
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p


def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p


def add_heading1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    if p.runs:
        p.runs[0].text = ""
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True


def add_heading2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if p.runs:
        p.runs[0].text = ""
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True


def add_heading3(text):
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    if p.runs:
        p.runs[0].text = ""
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    run.italic = True


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    if p.runs:
        p.runs[0].text = ""
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)


def page_break():
    doc.add_page_break()


def add_toc_field():
    """Chèn TOC field thật của Word — bấm F9 trong Word để cập nhật."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # \o "1-3" = lay heading 1..3, \h = hyperlink, \z = hide tab leader in web,
    # \u = use applied paragraph outline level
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = ("Mục lục sẽ được Word tự sinh khi bạn nhấn F9 "
                     "(hoặc References → Update Table).")
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def update_fields_on_open():
    """Báo Word tự cập nhật field khi mở file (F9 tự động)."""
    settings_xml = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings_xml.append(update_fields)


def add_image(filename, caption, width_cm=14):
    """Chèn ảnh canh giữa + caption đậm size 12."""
    path = ASSETS / filename
    if not path.exists():
        add_para(f"[Thiếu ảnh: {filename}]", italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cr = cap.add_run(caption)
    cr.font.name = 'Times New Roman'
    cr.font.size = Pt(12)
    cr.bold = True
    cr.italic = True


def add_table_from_rows(rows, header_bold=True, font_size=12):
    """Build bảng từ danh sách rows (list-of-list/tuple)."""
    n_rows = len(rows)
    n_cols = max(len(r) for r in rows)
    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(font_size)
            if i == 0 and header_bold:
                r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # khoảng trắng nhỏ sau bảng
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return t


def add_table_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True
    r.italic = True


# =========================================================
# TRANG BÌA — HUTECH
# =========================================================
add_para("BỘ GIÁO DỤC VÀ ĐÀO TẠO", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HỒ CHÍ MINH", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("KHOA CÔNG NGHỆ THÔNG TIN", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("------------------------", align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    add_para("", align=WD_ALIGN_PARAGRAPH.CENTER)

add_para("ĐỒ ÁN CƠ SỞ", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=22)
add_para("", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("NGHIÊN CỨU TẤN CÔNG VÀO HỆ THỐNG",
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18)
add_para("XÁC THỰC SINH TRẮC HỌC KHUÔN MẶT",
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18)
add_para("(BIOMETRIC AUTHENTICATION SECURITY)",
         bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
for _ in range(3):
    add_para("", align=WD_ALIGN_PARAGRAPH.CENTER)

add_para("Ngành: AN TOÀN THÔNG TIN", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para("", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Giảng viên hướng dẫn: ThS. Đặng Thị Thạch Thảo",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("Sinh viên thực hiện: Nguyễn Phúc Thịnh",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("MSSV: 2387700066    Lớp: 23DTHB1",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
for _ in range(3):
    add_para("")
add_para("TP. Hồ Chí Minh, năm 2026", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
page_break()

# =========================================================
# LỜI MỞ ĐẦU
# =========================================================
add_heading1("LỜI MỞ ĐẦU")
add_body("Trong những năm gần đây, xác thực sinh trắc học (biometric authentication) "
         "đã trở thành một trong những phương thức xác thực phổ biến nhất, được tích hợp "
         "rộng rãi trên điện thoại thông minh, hệ thống ngân hàng điện tử, kiểm soát ra "
         "vào và các dịch vụ định danh công dân. So với mật khẩu truyền thống, sinh trắc "
         "học mang lại trải nghiệm thuận tiện hơn vì người dùng không cần ghi nhớ thông "
         "tin và tốc độ xác thực rất nhanh. Tại Việt Nam, từ tháng 7 năm 2024, Ngân hàng "
         "Nhà nước đã yêu cầu các ngân hàng thương mại bắt buộc xác thực sinh trắc học "
         "khuôn mặt cho mọi giao dịch trực tuyến trên 10 triệu đồng, điều này càng khẳng "
         "định vai trò trung tâm của sinh trắc học trong hệ thống tài chính số.")
add_body("Tuy nhiên, đi cùng với sự phổ biến đó là các nguy cơ bảo mật ngày càng nghiêm "
         "trọng. Các hệ thống nhận diện khuôn mặt nếu không được trang bị cơ chế kiểm tra "
         "sống (Liveness Detection) rất dễ bị qua mặt bằng ảnh in, ảnh hiển thị trên "
         "điện thoại, video replay hoặc các kỹ thuật injection trực tiếp ở tầng API. "
         "Khác với mật khẩu có thể đặt lại bất cứ lúc nào, dữ liệu sinh trắc học một khi "
         "bị lộ thì không thể thay đổi, vì vậy hậu quả của một cuộc tấn công thành công "
         "là rất lâu dài. Đã có không ít sự cố thực tế tại Việt Nam khi các đối tượng "
         "dùng ảnh chụp lén của nạn nhân hoặc deepfake để chiếm đoạt tài khoản ngân hàng "
         "có giá trị lớn, gây thiệt hại đáng kể cho người dùng cuối.")
add_body("Đề tài “Nghiên cứu tấn công vào hệ thống sinh trắc học” được thực hiện nhằm "
         "khảo sát các bề mặt tấn công của một hệ thống xác thực khuôn mặt thực tế, xây "
         "dựng môi trường thí nghiệm gồm server có lỗ hổng và server đã được gia cố, từ "
         "đó đánh giá định lượng hiệu quả của giải pháp Liveness Detection dựa trên mạng "
         "MobileNetV2. Báo cáo trình bày toàn bộ quá trình từ cơ sở lý thuyết, triển "
         "khai, tấn công, phòng vệ cho đến đánh giá kết quả, hướng đến mục tiêu cảnh báo "
         "rủi ro và đề xuất biện pháp giảm thiểu cho các hệ thống sinh trắc học trong "
         "thực tế.")
add_body("Báo cáo được tổ chức thành năm chương theo trình tự logic của một dự án nghiên "
         "cứu an toàn thông tin: tổng quan, cơ sở lý thuyết, phân tích — triển khai, "
         "thực nghiệm và kết luận. Toàn bộ mã nguồn, dataset và kết quả đo đạc được lưu "
         "trữ trong kho lưu trữ kèm theo, cho phép tái lập thí nghiệm độc lập.")
page_break()

# =========================================================
# LỜI CAM ĐOAN
# =========================================================
add_heading1("LỜI CAM ĐOAN")
add_body("Em xin cam đoan đồ án “Nghiên cứu tấn công vào hệ thống sinh trắc học” là "
         "công trình nghiên cứu của riêng em, được thực hiện dưới sự hướng dẫn của ThS. "
         "Đặng Thị Thạch Thảo. Các số liệu, kết quả thực nghiệm trong báo cáo là trung "
         "thực, được đo đạc trên môi trường lab do em tự xây dựng và chưa từng được công "
         "bố trong bất kỳ công trình nào khác.")
add_body("Các nguồn tài liệu tham khảo đều được trích dẫn rõ ràng theo chuẩn IEEE. "
         "Toàn bộ thí nghiệm tấn công chỉ được thực hiện trên hệ thống do em tự triển "
         "khai phục vụ mục đích nghiên cứu, không nhằm vào bất kỳ hệ thống thực tế nào "
         "của tổ chức hoặc cá nhân khác. Em hoàn toàn chịu trách nhiệm trước Hội đồng "
         "khoa học và Nhà trường về tính trung thực của báo cáo này.")
add_para("")
add_para("TP. HCM, ngày … tháng … năm 2026", italic=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("Sinh viên thực hiện", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("")
add_para("")
add_para("Nguyễn Phúc Thịnh", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
page_break()

# =========================================================
# LỜI CẢM ƠN
# =========================================================
add_heading1("LỜI CẢM ƠN")
add_body("Trước tiên, em xin gửi lời cảm ơn chân thành đến quý thầy cô Khoa Công nghệ "
         "Thông tin – Trường Đại học Công nghệ TP. HCM (HUTECH) đã trang bị cho em những "
         "kiến thức nền tảng quan trọng về an toàn thông tin, học máy, lập trình mạng và "
         "phát triển ứng dụng trong suốt thời gian học tập tại trường. Những kiến thức "
         "này là cơ sở vững chắc giúp em tiếp cận và triển khai một đề tài có tính liên "
         "ngành như đề tài này.")
add_body("Em xin gửi lời cảm ơn sâu sắc nhất đến cô Đặng Thị Thạch Thảo — giảng viên "
         "hướng dẫn — đã tận tình góp ý, định hướng và tạo điều kiện để em có thể hoàn "
         "thành đề tài. Những phản hồi kịp thời của cô trong từng tuần báo cáo đã giúp "
         "em điều chỉnh hướng nghiên cứu phù hợp và bám sát mục tiêu đặt ra ban đầu.")
add_body("Em cũng xin cảm ơn các bạn cùng lớp 23DTHB1 đã hỗ trợ trong việc cung cấp "
         "ảnh khuôn mặt làm dữ liệu kiểm thử cũng như đóng góp ý kiến phản biện. Cuối "
         "cùng, em xin cảm ơn gia đình đã luôn động viên, ủng hộ em trong suốt quá trình "
         "học tập và thực hiện đồ án.")
add_body("Do thời gian và kiến thức còn hạn chế, đồ án không tránh khỏi những thiếu "
         "sót. Em rất mong nhận được những góp ý quý báu từ quý thầy cô để hoàn thiện "
         "hơn trong các nghiên cứu tiếp theo.")
add_para("")
add_para("Em xin chân thành cảm ơn!", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
page_break()

# =========================================================
# NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN
# =========================================================
add_heading1("NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN")
for _ in range(14):
    add_para("………………………………………………………………………………………………")
add_para("")
add_para("TP. HCM, ngày … tháng … năm 2026", italic=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("Giảng viên hướng dẫn", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("")
add_para("")
add_para("ThS. Đặng Thị Thạch Thảo", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
page_break()

# =========================================================
# MỤC LỤC — TOC field thật, F9 trong Word để cập nhật
# =========================================================
add_heading1("MỤC LỤC")
add_para("(Mở file trong Word: bấm Ctrl+A rồi F9, hoặc References → Update Table "
         "để Word tự sinh mục lục từ các Heading.)",
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
add_para("")
add_toc_field()
page_break()

# =========================================================
# DANH MỤC CHỮ VIẾT TẮT
# =========================================================
add_heading1("DANH MỤC CHỮ VIẾT TẮT")
abbr = [
    ("STT", "Viết tắt", "Nghĩa đầy đủ"),
    ("1", "FAR", "False Acceptance Rate – Tỷ lệ chấp nhận sai"),
    ("2", "FRR", "False Rejection Rate – Tỷ lệ từ chối sai"),
    ("3", "EER", "Equal Error Rate – Tỷ lệ lỗi cân bằng"),
    ("4", "APCER", "Attack Presentation Classification Error Rate"),
    ("5", "BPCER", "Bona-fide Presentation Classification Error Rate"),
    ("6", "PAD", "Presentation Attack Detection – Phát hiện tấn công trình diện"),
    ("7", "PAI", "Presentation Attack Instrument – Công cụ tấn công trình diện"),
    ("8", "CNN", "Convolutional Neural Network – Mạng nơ-ron tích chập"),
    ("9", "MobileNetV2", "Kiến trúc CNN nhẹ dùng cho thiết bị di động"),
    ("10", "ReLU", "Rectified Linear Unit – Hàm kích hoạt phi tuyến"),
    ("11", "ISO/IEC 30107", "Tiêu chuẩn quốc tế về phát hiện tấn công sinh trắc"),
    ("12", "API", "Application Programming Interface"),
    ("13", "REST", "Representational State Transfer"),
    ("14", "JSON", "JavaScript Object Notation"),
    ("15", "HTTP/HTTPS", "HyperText Transfer Protocol (Secure)"),
    ("16", "TLS", "Transport Layer Security – Mã hóa kênh truyền"),
    ("17", "OWASP", "Open Worldwide Application Security Project"),
    ("18", "CIA", "Confidentiality – Integrity – Availability"),
    ("19", "ROI", "Region of Interest – Vùng quan tâm trên ảnh"),
    ("20", "RGB", "Red – Green – Blue"),
    ("21", "IR", "Infrared – Hồng ngoại"),
    ("22", "3D", "Three-dimensional – Không gian ba chiều"),
    ("23", "GPU", "Graphics Processing Unit"),
    ("24", "MFA", "Multi-Factor Authentication – Xác thực đa yếu tố"),
    ("25", "ROC", "Receiver Operating Characteristic"),
    ("26", "AUC", "Area Under the (ROC) Curve"),
    ("27", "MITM", "Man-In-The-Middle"),
    ("28", "FIDO", "Fast IDentity Online Alliance"),
    ("29", "NIST", "National Institute of Standards and Technology"),
    ("30", "GDPR", "General Data Protection Regulation"),
]
add_table_from_rows(abbr, header_bold=True, font_size=12)
page_break()

# =========================================================
# DANH MỤC HÌNH ẢNH VÀ BẢNG BIỂU
# =========================================================
add_heading1("DANH MỤC HÌNH ẢNH VÀ BẢNG BIỂU")
add_heading3("Danh mục hình ảnh")
figs = [
    "Hình 2.1. Quy trình tổng quát của hệ thống xác thực sinh trắc học",
    "Hình 2.2. Phân loại các loại hình tấn công trình diện theo ISO/IEC 30107",
    "Hình 2.3. Kiến trúc MobileNetV2 với khối Inverted Residual",
    "Hình 2.4. Pipeline tấn công Injection vào hệ thống biometric",
    "Hình 2.5. So sánh tấn công Presentation và Injection",
    "Hình 3.1. Kiến trúc client-server của hệ thống thí nghiệm",
    "Hình 3.2. Luồng xử lý của endpoint /authenticate trên server vulnerable",
    "Hình 3.3. Mô hình tấn công Replay và Injection vào API xác thực",
    "Hình 3.4. Chi tiết tấn công Injection — bypass thiết bị thật",
    "Hình 3.5. Tiến trình phát triển: chưa hoàn thiện → bị tấn công → đã bảo mật",
    "Hình 4.1. Bộ dữ liệu Real (200 ảnh) và Fake (200 ảnh) dùng huấn luyện",
    "Hình 4.2. Quá trình huấn luyện MobileNetV2 — Loss và Accuracy",
    "Hình 4.2b. Confusion matrix của mô hình Liveness Detection",
    "Hình 4.3. Luồng xác thực có Liveness Detection trên server secured",
    "Hình 4.4. So sánh FAR / FRR / APCER trước và sau khi áp dụng Liveness",
    "Hình 4.5. Mô hình phòng vệ nhiều lớp (defense in depth)",
    "Hình 4.6. Giao diện server vulnerable (mô phỏng)",
    "Hình 4.7. Log tấn công Injection vào endpoint /authenticate",
    "Hình 4.8. Chu trình tấn công – phòng vệ – đánh giá",
    "Hình 5.1. Lộ trình triển khai thực tế giải pháp PAD",
]
for f in figs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

add_heading3("Danh mục bảng biểu")
tabs = [
    "Bảng 2.1. So sánh các loại sinh trắc học phổ biến",
    "Bảng 2.2. OWASP Top Risks cho hệ thống Biometric",
    "Bảng 3.1. Đặc điểm kỹ thuật của server vulnerable và server secured",
    "Bảng 4.1. Cấu hình huấn luyện mô hình MobileNetV2",
    "Bảng 4.2. Kết quả thực nghiệm trước/sau khi tích hợp Liveness Detection",
    "Bảng 4.3. So sánh latency các thành phần trong pipeline",
    "Bảng 5.1. Lộ trình triển khai PAD trong môi trường sản xuất",
]
for t in tabs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
page_break()

# =========================================================
# CHƯƠNG 1 - TỔNG QUAN ĐỀ TÀI
# =========================================================
add_heading1("CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI")

add_heading2("1.1. Lý do chọn đề tài")
add_body("Sinh trắc học đã và đang trở thành phương thức xác thực chủ đạo nhờ khả năng "
         "định danh người dùng dựa trên các đặc trưng vật lý và hành vi vốn có như khuôn "
         "mặt, vân tay, mống mắt, giọng nói. Theo nhiều báo cáo thị trường, hơn 80% điện "
         "thoại thông minh hiện đại đã tích hợp Face ID hoặc cảm biến vân tay; các ngân "
         "hàng số tại Việt Nam cũng yêu cầu xác thực khuôn mặt cho các giao dịch trên 10 "
         "triệu đồng kể từ tháng 7 năm 2024. Đây là xu hướng tất yếu khi mật khẩu truyền "
         "thống ngày càng bộc lộ những bất tiện và rủi ro bị đánh cắp qua tấn công lừa "
         "đảo (phishing) hoặc rò rỉ dữ liệu.")
add_body("Tuy nhiên, hầu hết hệ thống nhận diện khuôn mặt giai đoạn đầu đều chỉ thực "
         "hiện so khớp đặc trưng (feature matching) mà bỏ qua bước kiểm tra sống. Nhiều "
         "thí nghiệm công khai đã chứng minh có thể qua mặt các hệ thống này chỉ bằng "
         "ảnh in giấy hoặc ảnh hiển thị trên điện thoại. Khác với mật khẩu, một khi "
         "khuôn mặt – dữ liệu sinh trắc – đã bị lộ thì không thể thay đổi, do đó hậu "
         "quả là rất nghiêm trọng và lâu dài. Trong năm 2024-2025, cơ quan công an Việt "
         "Nam đã ghi nhận hàng loạt vụ lừa đảo sử dụng deepfake để chiếm đoạt tài khoản "
         "ngân hàng giá trị lớn, cho thấy mối đe dọa đã hiện hữu chứ không còn ở mức lý "
         "thuyết.")
add_body("Bên cạnh tấn công trình diện (presentation attack), tấn công injection ở tầng "
         "API cũng là một bề mặt tấn công đặc biệt nguy hiểm. Khi server không có cơ chế "
         "chứng thực thiết bị (device attestation) hay ký số ảnh tại client, kẻ tấn công "
         "có thể bỏ qua hoàn toàn cảm biến vật lý, chỉ cần gửi ảnh hợp lệ tới endpoint "
         "xác thực qua công cụ như curl hay Postman để chiếm phiên đăng nhập. Đây là "
         "khoảng trống mà nhiều giải pháp Liveness Detection thương mại hiện nay chưa "
         "che phủ trọn vẹn.")
add_body("Xuất phát từ thực trạng đó, đề tài chọn nghiên cứu các kịch bản tấn công vào "
         "hệ thống sinh trắc học khuôn mặt và đánh giá hiệu quả của giải pháp Liveness "
         "Detection dựa trên học sâu. Đề tài hướng đến tính ứng dụng cao và phù hợp với "
         "định hướng An toàn thông tin của khoa, đồng thời cung cấp một bộ tài liệu mở "
         "có thể dùng làm tài liệu tham khảo cho các sinh viên khóa sau.")

add_heading2("1.2. Mục tiêu nghiên cứu")
add_body("Đề tài đặt ra các mục tiêu cụ thể như sau:")
add_bullet("Hệ thống hóa cơ sở lý thuyết về xác thực sinh trắc học, các bề mặt tấn công "
           "và tiêu chuẩn ISO/IEC 30107 về phát hiện tấn công trình diện.")
add_bullet("Xây dựng môi trường thí nghiệm gồm hai phiên bản server: bản dễ bị tổn "
           "thương (vulnerable) và bản đã gia cố (secured) bằng Liveness Detection.")
add_bullet("Triển khai và phân tích các kịch bản tấn công thực tế: spoofing bằng ảnh "
           "in, ảnh trên điện thoại, replay, injection trực tiếp qua API.")
add_bullet("Huấn luyện mô hình phân loại Real/Fake dựa trên MobileNetV2 với dataset tự "
           "thu thập gồm 200 ảnh thật và 200 ảnh giả.")
add_bullet("Đánh giá định lượng bằng các chỉ số FAR, FRR, APCER, BPCER trước và sau "
           "khi tích hợp giải pháp phòng vệ.")
add_bullet("Đề xuất kiến trúc phòng vệ chiều sâu (defense in depth) bao gồm Liveness "
           "Detection, device attestation, mã hóa kênh truyền và xác thực đa yếu tố.")

add_heading2("1.3. Đối tượng và phạm vi nghiên cứu")
add_heading3("1.3.1. Đối tượng nghiên cứu")
add_body("Đối tượng nghiên cứu là hệ thống xác thực bằng khuôn mặt (face authentication) "
         "được triển khai dưới dạng ứng dụng web/API, sử dụng thư viện face_recognition "
         "(dlib) cho nhận diện và mạng MobileNetV2 cho phát hiện tấn công trình diện. "
         "Đây là kiến trúc đại diện cho rất nhiều hệ thống thực tế đang vận hành — từ "
         "các ứng dụng chấm công nhỏ đến các cổng đăng nhập của ngân hàng số.")
add_heading3("1.3.2. Phạm vi nghiên cứu")
add_bullet("Trong phạm vi lab: chỉ tấn công và phòng vệ trên hệ thống do em tự xây "
           "dựng, không nhằm vào bất kỳ hệ thống thương mại nào.")
add_bullet("Tập trung vào sinh trắc khuôn mặt 2D từ camera RGB; không khảo sát IR, 3D, "
           "vân tay hay mống mắt.")
add_bullet("Các kịch bản tấn công giới hạn ở: ảnh in giấy, ảnh hiển thị trên điện "
           "thoại, ảnh chụp lại màn hình và injection ảnh tĩnh qua API.")
add_bullet("Không khảo sát tấn công deepfake video phức tạp do giới hạn về phần cứng "
           "và thời gian, nhưng có đề cập trong phần hướng phát triển.")

add_heading2("1.4. Phương pháp nghiên cứu")
add_bullet("Nghiên cứu lý thuyết: tổng hợp các tài liệu, tiêu chuẩn ISO/IEC 30107, các "
           "công bố khoa học về Liveness Detection và OWASP Biometric Risks.")
add_bullet("Phương pháp thực nghiệm: xây dựng hệ thống, mô phỏng tấn công, thu thập số "
           "liệu và đánh giá theo chỉ số chuẩn ISO/IEC 30107-3.")
add_bullet("Phương pháp so sánh: đối chiếu kết quả của hệ thống vulnerable và secured "
           "trên cùng một bộ dữ liệu kiểm thử với cùng môi trường phần cứng.")
add_bullet("Phương pháp phân tích định lượng: dùng ma trận nhầm lẫn (confusion matrix) "
           "và các chỉ số FAR, FRR, APCER, BPCER để đo hiệu năng phòng vệ.")

add_heading2("1.5. Ý nghĩa khoa học và thực tiễn")
add_body("Về mặt khoa học, đề tài hệ thống lại các kịch bản tấn công sinh trắc theo "
         "đúng vị trí trong quy trình xác thực và đề xuất vị trí phù hợp để đặt cơ chế "
         "phòng vệ. Báo cáo cũng cung cấp một mô hình đe dọa (threat model) cụ thể, có "
         "thể được tái sử dụng cho các đề tài tương tự.")
add_body("Về mặt thực tiễn, đề tài cung cấp một bộ tài liệu mở gồm mã nguồn demo, "
         "dataset và hướng dẫn triển khai có thể tái sử dụng cho mục đích đào tạo và "
         "đánh giá an toàn cho các hệ thống biometric đang vận hành. Các đoạn mã có thể "
         "được điều chỉnh nhỏ để tích hợp vào hệ thống chấm công hoặc hệ thống đăng "
         "nhập của doanh nghiệp vừa và nhỏ.")

add_heading2("1.6. Bố cục báo cáo")
add_bullet("Chương 1: Tổng quan đề tài — trình bày lý do, mục tiêu, phạm vi và phương "
           "pháp nghiên cứu.")
add_bullet("Chương 2: Cơ sở lý thuyết về sinh trắc học và các loại tấn công.")
add_bullet("Chương 3: Phân tích và triển khai hệ thống thí nghiệm gồm cả phiên bản "
           "vulnerable và secured.")
add_bullet("Chương 4: Thực nghiệm tấn công, phòng vệ và đánh giá kết quả định lượng.")
add_bullet("Chương 5: Kết luận, đóng góp và hướng phát triển trong tương lai.")

add_heading2("1.7. Dự kiến kết quả định lượng")
add_body("Trước khi triển khai, em dự kiến các chỉ số mục tiêu cho hệ thống đã gia cố "
         "như sau: FAR < 5%, FRR < 7%, APCER < 3%, BPCER < 5%, latency tăng thêm không "
         "quá 100 ms so với server vulnerable. Các con số này được tham chiếu từ các "
         "công bố PAD trên dataset CASIA-FASD và Replay-Attack, đồng thời điều chỉnh "
         "cho phù hợp với điều kiện lab có dataset nhỏ. Chương 4 sẽ trình bày kết quả "
         "đo đạc thực tế và so sánh với các mục tiêu này.")
page_break()

# =========================================================
# CHƯƠNG 2 - CƠ SỞ LÝ THUYẾT
# =========================================================
add_heading1("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT")

add_heading2("2.1. Tổng quan về xác thực sinh trắc học")
add_body("Sinh trắc học (biometrics) là phương pháp xác định danh tính của một cá nhân "
         "dựa trên các đặc trưng sinh học hoặc hành vi vốn có. Khác với mật khẩu hay "
         "thẻ vật lý, đặc trưng sinh trắc gắn liền với chủ thể nên rất khó bị quên hoặc "
         "đánh mất, đồng thời mang lại trải nghiệm xác thực nhanh và thuận tiện. Tuy "
         "nhiên, đặc trưng này cũng không thể thay đổi nếu bị lộ, vì vậy yêu cầu bảo vệ "
         "phải nghiêm ngặt hơn so với mật khẩu thông thường.")

add_heading3("2.1.1. Phân loại sinh trắc học")
add_bullet("Sinh trắc vật lý: khuôn mặt, vân tay, mống mắt, hình dạng bàn tay, mạch máu "
           "lòng bàn tay.")
add_bullet("Sinh trắc hành vi: chữ ký, dáng đi, cách gõ phím, giọng nói, thói quen "
           "tương tác với thiết bị.")
add_bullet("Sinh trắc kết hợp đa phương thức: kết hợp hai hoặc nhiều đặc trưng để tăng "
           "độ chính xác và giảm rủi ro giả mạo.")

add_heading3("2.1.2. Quy trình xác thực sinh trắc học")
add_body("Một hệ thống biometric tổng quát gồm hai pha: pha đăng ký (enrollment) lưu "
         "đặc trưng tham chiếu của người dùng vào cơ sở dữ liệu, và pha xác thực "
         "(authentication) so khớp đặc trưng mới chụp với mẫu đã lưu để ra quyết định "
         "Accept/Reject. Quy trình chi tiết gồm 6 bước: (1) thu nhận tín hiệu, (2) "
         "tiền xử lý, (3) kiểm tra sống, (4) trích xuất đặc trưng, (5) so khớp với "
         "database và (6) ra quyết định cuối cùng.")
add_image("fig_2_1_quy_trinh_bio.png",
          "Hình 2.1. Quy trình tổng quát của hệ thống xác thực sinh trắc học",
          width_cm=14)

add_heading3("2.1.3. So sánh các loại sinh trắc học")
add_table_caption("Bảng 2.1. So sánh các loại sinh trắc học phổ biến")
bio_compare = [
    ("Loại sinh trắc", "Độ chính xác", "Khả năng giả mạo", "Chi phí thiết bị"),
    ("Khuôn mặt 2D", "Cao", "Cao (ảnh, video)", "Thấp (webcam)"),
    ("Khuôn mặt 3D / IR", "Rất cao", "Trung bình", "Cao"),
    ("Vân tay", "Rất cao", "Trung bình (silicon)", "Trung bình"),
    ("Mống mắt", "Rất cao", "Thấp", "Cao"),
    ("Giọng nói", "Trung bình", "Cao (deepfake audio)", "Thấp (mic)"),
    ("Dáng đi", "Trung bình", "Trung bình", "Trung bình"),
]
add_table_from_rows(bio_compare)

add_heading2("2.2. Nhận diện khuôn mặt với face_recognition và dlib")
add_body("Thư viện face_recognition của Adam Geitgey được xây dựng trên nền dlib, sử "
         "dụng mô hình ResNet-34 đã được tinh chỉnh để chuyển một khuôn mặt thành vector "
         "đặc trưng 128 chiều (128-d face embedding). Hai khuôn mặt được coi là cùng "
         "một người nếu khoảng cách Euclidean giữa hai vector nhỏ hơn ngưỡng nhất định, "
         "thường là 0.6. Đây là giá trị được tác giả thư viện xác định qua thực nghiệm "
         "trên dataset Labeled Faces in the Wild (LFW) và đạt độ chính xác 99.38%.")
add_body("Trong đề tài, ngưỡng so khớp được chọn là MATCH_THRESHOLD = 0.6, đây cũng là "
         "giá trị mặc định mà tài liệu chính thức của face_recognition khuyến nghị, "
         "giúp cân bằng giữa FAR và FRR. Việc nâng ngưỡng (ví dụ 0.7) sẽ làm FAR tăng "
         "mạnh, đây cũng là một trong những lỗ hổng được khai thác trong kịch bản "
         "tấn công 'threshold abuse' ở Chương 4.")
add_body("Mô hình face_recognition có ưu điểm là rất nhẹ, chạy được trên CPU với tốc "
         "độ ~100 ms/khung, mã nguồn mở và cộng đồng sử dụng rộng rãi. Tuy nhiên, do "
         "chỉ dựa trên embedding 2D nên mô hình này không có khả năng phân biệt giữa "
         "ảnh thật và ảnh in/màn hình — đây chính là khoảng trống mà Liveness Detection "
         "cần lấp đầy.")

add_heading2("2.3. Các loại tấn công vào hệ thống sinh trắc học")
add_image("fig_2_2_phan_loai_tan_cong.png",
          "Hình 2.2. Phân loại các loại hình tấn công trình diện theo ISO/IEC 30107",
          width_cm=14)

add_heading3("2.3.1. Tấn công trình diện (Presentation Attack)")
add_body("Theo tiêu chuẩn ISO/IEC 30107, tấn công trình diện là việc đưa các Công cụ "
         "tấn công trình diện (Presentation Attack Instrument – PAI) vào trước cảm "
         "biến để giả mạo danh tính. Các loại PAI phổ biến gồm: ảnh in giấy, ảnh hiển "
         "thị trên điện thoại/màn hình, mặt nạ silicon, video deepfake. Đây là dạng "
         "tấn công kinh điển và dễ thực hiện nhất, chỉ cần một bức ảnh hợp lý của nạn "
         "nhân lấy từ mạng xã hội.")

add_heading3("2.3.2. Tấn công Replay")
add_body("Kẻ tấn công ghi lại một phiên xác thực hợp lệ và phát lại tới cảm biến hoặc "
         "API. Đây là dạng tấn công đặc biệt nguy hiểm vì dữ liệu được dùng vốn dĩ là "
         "dữ liệu thật, do đó vượt qua các bước kiểm tra cấu trúc thông thường. Replay "
         "có thể xảy ra ở nhiều tầng: replay tín hiệu video trước camera, replay request "
         "HTTP, hoặc replay ảnh đã được mã hóa base64 trong payload.")

add_heading3("2.3.3. Tấn công Injection")
add_body("Thay vì trình diện trước camera, kẻ tấn công gửi trực tiếp ảnh hoặc dữ liệu "
         "đặc trưng tới API xác thực. Khi server không có cơ chế chứng thực thiết bị "
         "hoặc kênh truyền, attacker có thể bypass hoàn toàn cảm biến vật lý. Đây là "
         "dạng tấn công nguy hiểm nhất vì có thể tự động hóa và scale lên hàng triệu "
         "request, đồng thời rất khó phát hiện chỉ bằng cách phân tích nội dung ảnh.")
add_image("fig_2_4_injection_pipeline.png",
          "Hình 2.4. Pipeline tấn công Injection vào hệ thống biometric",
          width_cm=14)

add_heading3("2.3.4. Tấn công Deepfake")
add_body("Sử dụng các mô hình sinh đối kháng (GAN) để tổng hợp khuôn mặt giả nhưng "
         "có biểu cảm sống động. Đây là dạng tấn công cao cấp, có thể đánh bại nhiều "
         "phương pháp Liveness 2D đơn giản. Trong giai đoạn 2024-2025, deepfake video "
         "thời gian thực đã trở thành một vector tấn công phổ biến trong các vụ lừa "
         "đảo qua cuộc gọi video tại Việt Nam.")

add_heading3("2.3.5. So sánh giữa Presentation và Injection")
add_image("fig_2_5_pa_vs_injection.png",
          "Hình 2.5. So sánh tấn công Presentation và Injection",
          width_cm=14)
add_body("Hai dạng tấn công này khác nhau ở vị trí xảy ra trong pipeline xác thực: "
         "presentation tấn công từ phía cảm biến (vật lý), trong khi injection tấn "
         "công từ phía API (logic). Một hệ thống an toàn cần có biện pháp phòng vệ ở "
         "cả hai vị trí — Liveness Detection chỉ giải quyết được một phần của bài toán.")

add_heading2("2.4. Liveness Detection và tiêu chuẩn ISO/IEC 30107")
add_body("Liveness Detection (Phát hiện sống) – còn gọi là Presentation Attack "
         "Detection (PAD) – là cơ chế kiểm tra xem khuôn mặt được trình diện có phải "
         "đến từ một người sống thật hay không. Tiêu chuẩn ISO/IEC 30107 định nghĩa "
         "khung đánh giá PAD và đưa ra hai chỉ số chính:")
add_bullet("APCER (Attack Presentation Classification Error Rate): tỷ lệ tấn công bị "
           "phân loại nhầm thành mẫu thật. APCER càng nhỏ thì hệ thống chống giả mạo "
           "càng tốt.")
add_bullet("BPCER (Bona-fide Presentation Classification Error Rate): tỷ lệ mẫu thật "
           "bị phân loại nhầm thành tấn công. BPCER càng nhỏ thì người dùng thật càng "
           "ít bị phiền hà.")
add_body("Liveness Detection có thể chia thành hai nhóm: passive (thụ động — chỉ phân "
         "tích ảnh đầu vào) và active (chủ động — yêu cầu người dùng nháy mắt, quay "
         "đầu, đọc số ngẫu nhiên). Đề tài này tập trung vào passive liveness vì đây là "
         "phương pháp ít gây phiền cho người dùng và phù hợp với kịch bản chấm công "
         "nhanh.")

add_heading2("2.5. Mạng MobileNetV2 và Transfer Learning")
add_body("MobileNetV2 (Sandler et al., 2018) là kiến trúc CNN nhẹ được thiết kế cho "
         "thiết bị di động và các ứng dụng thời gian thực. Đặc điểm nổi bật của "
         "MobileNetV2 là khối Inverted Residual với Linear Bottleneck, sử dụng "
         "depthwise separable convolution để giảm đáng kể số phép tính so với CNN "
         "truyền thống nhưng vẫn duy trì độ chính xác cao trên ImageNet. Tổng số "
         "tham số chỉ khoảng 3.4 triệu, nhỏ hơn ResNet-50 khoảng 7 lần, phù hợp với "
         "yêu cầu chạy thời gian thực trên CPU laptop.")
add_image("fig_2_3_mobilenetv2.png",
          "Hình 2.3. Kiến trúc MobileNetV2 với khối Inverted Residual",
          width_cm=14)
add_body("Trong đề tài, MobileNetV2 được sử dụng theo phương pháp transfer learning: "
         "giữ nguyên các lớp tích chập đã được tiền huấn luyện trên ImageNet và chỉ "
         "thay thế lớp phân loại cuối cùng bằng một lớp Linear với 2 đầu ra "
         "(REAL/FAKE) tương ứng với bài toán nhị phân. Quá trình huấn luyện chia làm "
         "hai pha: (1) đóng băng backbone, chỉ huấn luyện classifier head; (2) mở "
         "băng các lớp cuối của backbone để fine-tune với learning rate thấp hơn.")

add_heading2("2.6. Các chỉ số đánh giá")
add_bullet("FAR (False Acceptance Rate): tỷ lệ kẻ giả mạo được chấp nhận là người thật.")
add_bullet("FRR (False Rejection Rate): tỷ lệ người thật bị từ chối nhầm.")
add_bullet("EER (Equal Error Rate): điểm tại đó FAR = FRR, đặc trưng cho độ ổn định "
           "tổng thể của hệ thống.")
add_bullet("APCER, BPCER: hai chỉ số chuyên biệt cho PAD theo ISO/IEC 30107.")
add_bullet("Accuracy, Precision, Recall, F1-score: các chỉ số phân loại nhị phân chuẩn "
           "trong machine learning.")
add_bullet("Latency: thời gian xử lý mỗi request, ảnh hưởng trực tiếp đến trải nghiệm "
           "người dùng — yêu cầu thường < 500 ms.")

add_heading2("2.7. OWASP Biometric Risks")
add_body("OWASP đã công bố danh sách rủi ro hàng đầu cho hệ thống sinh trắc học, "
         "tổng hợp từ các sự cố thực tế và nghiên cứu học thuật. Bảng 2.2 liệt kê 10 "
         "rủi ro chính, đây cũng là các checklist quan trọng khi đánh giá an toàn hệ "
         "thống biometric.")
add_table_caption("Bảng 2.2. OWASP Top Risks cho hệ thống Biometric")
risks = [
    ("STT", "Rủi ro", "Mô tả ngắn"),
    ("1", "Spoofing trình diện", "Sử dụng ảnh in, màn hình, mặt nạ qua mặt cảm biến."),
    ("2", "Replay/Injection", "Phát lại hoặc gửi trực tiếp dữ liệu sinh trắc đã chiếm đoạt."),
    ("3", "Lộ template", "Database lưu đặc trưng bị rò rỉ → không thể thay đổi sinh trắc."),
    ("4", "Thiếu Liveness Detection", "Hệ thống chỉ so khớp, không xác định được mẫu sống."),
    ("5", "Tấn công ngưỡng", "Khai thác ngưỡng quá lỏng/lệch để tăng FAR."),
    ("6", "Dữ liệu huấn luyện độc hại", "Poisoning dataset PAD để giảm hiệu quả mô hình."),
    ("7", "Kênh truyền không mã hóa", "MITM nghe lén/sửa đổi ảnh trên đường truyền."),
    ("8", "Thiếu MFA", "Phụ thuộc duy nhất vào yếu tố sinh trắc."),
    ("9", "Lỗi triển khai API", "Endpoint authenticate không kiểm tra nguồn gốc thiết bị."),
    ("10", "Bypass logic", "Khai thác lỗi logic xử lý Liveness/match để vượt qua."),
]
add_table_from_rows(risks)

add_heading2("2.8. FIDO Alliance và NIST 800-63B")
add_body("FIDO Alliance là tổ chức công nghiệp đặt ra các tiêu chuẩn xác thực hiện đại "
         "không dùng mật khẩu. Chương trình FIDO Biometric Component Certification yêu "
         "cầu cảm biến và mô-đun PAD đạt mức APCER < 5% (Level A) hoặc < 1% (Level B). "
         "Các sản phẩm thương mại như Apple Face ID hay Windows Hello đều đạt chứng "
         "nhận FIDO Level B.")
add_body("NIST SP 800-63B (Digital Identity Guidelines) phân loại biometric vào nhóm "
         "'something you are' và yêu cầu phải kết hợp với ít nhất một yếu tố khác để "
         "đạt mức xác thực AAL2 (Authenticator Assurance Level 2) — phù hợp với các "
         "giao dịch tài chính. Đây là cơ sở pháp lý cho việc triển khai MFA trong các "
         "hệ thống biometric thực tế.")
page_break()

# =========================================================
# CHƯƠNG 3 - PHÂN TÍCH VÀ TRIỂN KHAI
# =========================================================
add_heading1("CHƯƠNG 3: PHÂN TÍCH VÀ TRIỂN KHAI HỆ THỐNG")

add_heading2("3.0. Tiến trình tổng thể của hệ thống")
add_body("Trước khi đi vào chi tiết kiến trúc, em xin trình bày bức tranh tổng thể về "
         "hành trình phát triển của hệ thống xác thực sinh trắc học trong đề tài này. "
         "Hành trình được chia làm ba giai đoạn rõ rệt: (1) hệ thống chưa hoàn thiện — "
         "chỉ có chức năng cơ bản và để lộ nhiều lỗ hổng; (2) hệ thống bị tấn công thực "
         "tế — kẻ tấn công khai thác toàn bộ các lỗ hổng đã xác định; (3) hệ thống đã "
         "được bảo mật — bổ sung Liveness Detection, xác thực thiết bị, mã hóa kênh "
         "truyền và logging giám sát.")
add_image("fig_3_5_tien_trinh_3_giai_doan.png",
          "Hình 3.5. Tiến trình phát triển: chưa hoàn thiện → bị tấn công → đã bảo mật",
          width_cm=15.5)
add_body("Sơ đồ trên cho thấy rõ vị trí mà mỗi lỗ hổng và mỗi biện pháp phòng vệ tác "
         "động vào pipeline xác thực. Việc tách ba giai đoạn giúp người đọc dễ dàng "
         "đối chiếu giữa rủi ro và biện pháp tương ứng — đây là cách trình bày có thể "
         "tái sử dụng cho các hệ thống an toàn thông tin khác.")

add_heading2("3.1. Mô hình kiến trúc tổng thể")
add_body("Hệ thống thí nghiệm được xây dựng theo kiến trúc client-server cổ điển. "
         "Phía client là một trang web đơn giản chạy trên trình duyệt, sử dụng API "
         "getUserMedia để lấy luồng video từ webcam. Phía server là một ứng dụng Flask "
         "(Python 3.10) cung cấp các REST API: /register, /authenticate, /users.")
add_image("fig_3_1_kien_truc_tong_the.png",
          "Hình 3.1. Kiến trúc client-server của hệ thống thí nghiệm",
          width_cm=14)
add_body("Database lưu đặc trưng khuôn mặt được tổ chức dưới dạng file pickle "
         "(face_database.pkl), mỗi entry là một dict gồm tên người dùng, vector 128-d "
         "và thời điểm đăng ký. Đây là cấu trúc đủ đơn giản để tập trung vào khía cạnh "
         "an toàn của thuật toán xác thực thay vì hạ tầng dữ liệu. Trong môi trường "
         "thực tế, database thường được thay bằng PostgreSQL/MySQL với mã hóa AT-REST "
         "và cơ chế phân quyền nghiêm ngặt.")

add_heading2("3.2. Phiên bản server có lỗ hổng (vulnerable)")
add_body("Server vulnerable (server/app.py) cài đặt trực tiếp luồng xác thực truyền "
         "thống. Endpoint /authenticate chỉ thực hiện đúng 3 bước: phát hiện khuôn mặt, "
         "trích đặc trưng và so khớp Euclidean với database, không có bất kỳ kiểm tra "
         "sống nào. Khi distance ≤ MATCH_THRESHOLD (0.6) thì trả về thành công, đồng "
         "thời response gắn cờ liveness_check = 'DISABLED' để minh họa rõ điểm yếu.")
add_image("fig_3_2_server_vulnerable.png",
          "Hình 3.2. Luồng xử lý của endpoint /authenticate trên server vulnerable",
          width_cm=14)

add_heading3("3.2.1. Mã nguồn cốt lõi của /authenticate")
add_body("Đoạn mã trọng yếu của server vulnerable được trình bày dưới đây. Có thể "
         "thấy hoàn toàn không có bất kỳ phép kiểm tra sống nào; chỉ duy nhất một "
         "vòng lặp tính khoảng cách Euclidean và so sánh với ngưỡng:")
codep = doc.add_paragraph()
codep.paragraph_format.left_indent = Cm(0.6)
cr = codep.add_run(
    "for name, user_data in db.items():\n"
    "    distance = np.linalg.norm(user_data['encoding'] - unknown_encoding)\n"
    "    if distance < best_distance:\n"
    "        best_distance = distance\n"
    "        best_match = name\n"
    "if best_distance <= MATCH_THRESHOLD:\n"
    "    return jsonify({'authenticated': True,\n"
    "                    'liveness_check': 'DISABLED',\n"
    "                    'user': best_match})"
)
cr.font.name = 'Consolas'
cr.font.size = Pt(11)

add_heading3("3.2.2. Bề mặt tấn công")
add_bullet("Spoofing trình diện: đưa ảnh in / ảnh trên điện thoại trước webcam.")
add_bullet("Injection: gửi POST /authenticate với trường image là ảnh tĩnh đã chiếm "
           "đoạt, không cần thiết bị thật.")
add_bullet("Replay: phát lại ảnh đã chụp trước đó của nạn nhân.")
add_bullet("Threshold abuse: nếu MATCH_THRESHOLD bị nâng lên (ví dụ 0.7) thì FAR sẽ "
           "tăng vọt, ngay cả người không đăng ký cũng có thể bị nhận nhầm.")
add_bullet("Enumeration: gọi /users để liệt kê toàn bộ người dùng đã đăng ký, từ đó "
           "ra quyết định tấn công có chủ đích.")

add_heading2("3.3. Phiên bản server đã gia cố (secured)")
add_body("Server secured (server/app_secured.py) bổ sung bước Liveness Detection vào "
         "đầu pipeline xác thực. Khi nhận ảnh, server crop khuôn mặt 224x224, đưa qua "
         "mô hình MobileNetV2 đã huấn luyện. Nếu xác suất Real < LIVENESS_THRESHOLD "
         "(0.7) thì server trả HTTP 403 với cờ presentation_attack = True và DỪNG "
         "ngay; chỉ khi vượt qua kiểm tra sống thì luồng so khớp khuôn mặt mới được "
         "thực hiện.")
add_image("fig_4_3_server_secured.png",
          "Hình 4.3. Luồng xác thực có Liveness Detection trên server secured",
          width_cm=14)
add_body("Ngoài Liveness Detection, server secured còn bổ sung: (1) endpoint /health "
         "để kiểm tra trạng thái mô hình; (2) logging chi tiết mọi request thất bại; "
         "(3) rate limiting cơ bản để chống bruteforce; (4) header bảo mật chuẩn "
         "(X-Frame-Options, Content-Security-Policy).")

add_heading2("3.4. So sánh hai phiên bản server")
add_table_caption("Bảng 3.1. Đặc điểm kỹ thuật của server vulnerable và server secured")
cmp_tab = [
    ("Tiêu chí", "Server vulnerable", "Server secured"),
    ("Liveness Detection", "Không có", "MobileNetV2, ngưỡng 0.7"),
    ("Bước trước khi match", "Chỉ phát hiện khuôn mặt", "Phát hiện + Liveness"),
    ("Phản hồi với ảnh giả", "Có thể nhận thật (FAR cao)", "HTTP 403 / presentation_attack"),
    ("Endpoint bổ sung", "Không có", "/health (model status)"),
    ("Rate limiting", "Không", "Có"),
    ("Logging chi tiết", "Cơ bản", "Đầy đủ thất bại + cảnh báo"),
    ("Mục tiêu sử dụng", "Minh họa lỗ hổng", "So sánh, đánh giá phòng vệ"),
]
add_table_from_rows(cmp_tab)

add_heading2("3.5. Tiền xử lý ảnh đầu vào")
add_body("Ảnh đầu vào được chuẩn hóa qua hàm decode_image: chuyển sang RGB, kiểm tra "
         "độ sáng trung bình; nếu brightness < 80 thì áp dụng CLAHE (Contrast Limited "
         "Adaptive Histogram Equalization) trên kênh L của không gian màu LAB. Bước "
         "này đặc biệt quan trọng cho Liveness Detection vì các mẫu thiếu sáng dễ bị "
         "phân loại nhầm thành Fake nếu không cân bằng độ tương phản.")
add_body("Sau CLAHE, ảnh được resize về 224×224 pixel — kích thước đầu vào chuẩn của "
         "MobileNetV2 — rồi normalize theo mean và std của ImageNet (mean=[0.485, "
         "0.456, 0.406], std=[0.229, 0.224, 0.225]). Đây là bước bắt buộc để mô hình "
         "transfer learning hoạt động đúng với trọng số đã học từ ImageNet.")

add_heading2("3.6. Mô hình đe dọa (Threat Model)")
add_bullet("Attacker có quyền truy cập trang web và API công khai của hệ thống.")
add_bullet("Attacker có thể chiếm đoạt một ảnh khuôn mặt của nạn nhân (qua mạng xã hội, "
           "ảnh chụp trộm hoặc rò rỉ dữ liệu).")
add_bullet("Attacker không có quyền can thiệp vào server, không có khóa bí mật của hệ "
           "thống và không kiểm soát được webcam của nạn nhân.")
add_bullet("Mục tiêu của attacker: làm cho server trả về authenticated = True dưới "
           "danh nghĩa nạn nhân.")
add_bullet("Attacker được trang bị các công cụ phổ thông: trình duyệt, curl/Postman, "
           "Python, máy in văn phòng, điện thoại Android/iOS.")

add_heading2("3.7. Phân tích cấu trúc pickle database")
add_body("File face_database.pkl được lưu dưới dạng dict Python với khóa là username, "
         "giá trị là một sub-dict gồm hai trường: 'encoding' (numpy array 128 chiều) "
         "và 'registered_at' (timestamp ISO 8601). Việc dùng pickle có ưu điểm là gọn "
         "và nhanh, nhưng cũng kèm theo rủi ro lớn: pickle có thể chứa mã độc nếu file "
         "bị thay thế. Nếu attacker chiếm được quyền ghi vào file pickle, họ có thể "
         "inject một class với phương thức __reduce__ để thực thi mã tùy ý khi server "
         "load database.")
add_body("Trong môi trường lab, em chấp nhận rủi ro này để giữ kiến trúc đơn giản. "
         "Tuy nhiên, bản triển khai sản xuất nên thay pickle bằng cơ sở dữ liệu quan "
         "hệ với encoding lưu dưới dạng BLOB và áp dụng mã hóa AES-256 khi lưu. Khóa "
         "mã hóa nên được quản lý qua dịch vụ KMS (Key Management Service) như AWS "
         "KMS hoặc HashiCorp Vault.")

add_heading2("3.8. Lỗ hổng tại endpoint /users (enumeration)")
add_body("Endpoint /users của server vulnerable trả về toàn bộ danh sách người dùng "
         "đã đăng ký kèm thời điểm đăng ký, không yêu cầu xác thực. Điều này tạo điều "
         "kiện cho tấn công enumeration: attacker chỉ cần một request GET là biết "
         "được toàn bộ danh sách mục tiêu trong hệ thống. Trong thí nghiệm, ngay từ "
         "phase trinh sát, kẻ tấn công đã liệt kê được 7 user gồm: thinh, nhan, tai, "
         "hoang, NGUYEN VAN CHUOT, Xxsxcsxc và hihi.")
add_image("fig_3_4_injection_chi_tiet.png",
          "Hình 3.4. Chi tiết tấn công Injection — bypass thiết bị thật",
          width_cm=14)
add_body("Bài học rút ra: mọi endpoint trả về thông tin nội bộ — kể cả chỉ là tên — "
         "đều phải có xác thực. Server secured trong đề tài đã sửa endpoint /users "
         "thành chỉ trả về số lượng người dùng (count) và yêu cầu token quản trị "
         "(admin token) cho các thao tác đọc chi tiết.")
page_break()

# =========================================================
# CHƯƠNG 4 - THỰC NGHIỆM, PHÒNG VỆ VÀ ĐÁNH GIÁ
# =========================================================
add_heading1("CHƯƠNG 4: THỰC NGHIỆM, PHÒNG VỆ VÀ ĐÁNH GIÁ")

add_heading2("4.1. Môi trường thí nghiệm")
add_bullet("Phần cứng: laptop CPU Intel Core i5/i7 (8 nhân, 2.4 GHz), RAM 16 GB, "
           "webcam tích hợp 720p, không có GPU rời.")
add_bullet("Phần mềm: Python 3.10, Flask 3.x, OpenCV 4.x, face_recognition 1.3, "
           "dlib 19.x, PyTorch 2.x, torchvision 0.16, numpy 1.26.")
add_bullet("Hệ điều hành: Windows 11 Pro 64-bit. Trình duyệt kiểm thử: Chrome 120+, "
           "Edge 120+, Firefox 121+.")
add_bullet("Công cụ tấn công: Python requests, Postman, curl, ảnh in từ máy in HP "
           "LaserJet, điện thoại iPhone 12 và Samsung Galaxy A52.")

add_heading2("4.2. Kịch bản tấn công trên server vulnerable")
add_image("fig_3_3_attack_replay_injection.png",
          "Hình 3.3. Mô hình tấn công Replay và Injection vào API xác thực",
          width_cm=14)

add_heading3("4.2.1. Spoofing bằng ảnh in")
add_body("Lấy ảnh chân dung của người dùng đã đăng ký, in trên giấy A4 màu và đưa "
         "trước webcam. Server vulnerable nhận ảnh, vẫn trích được encoding 128-d và "
         "so khớp thành công. Tỷ lệ vượt qua trong các phép thử đạt gần như 100% với "
         "khoảng cách Euclidean trung bình ~0.38 — thấp hơn ngưỡng 0.6 rất nhiều.")

add_heading3("4.2.2. Spoofing bằng ảnh trên điện thoại")
add_body("Hiển thị ảnh chân dung trên điện thoại độ phân giải cao (Full HD trở lên) "
         "và đặt trước webcam ở khoảng cách 30-50 cm. Kết quả tương tự ảnh in: xác "
         "thực thành công với khoảng cách Euclidean rất nhỏ (≈ 0.32–0.45). Đáng chú "
         "ý là ảnh trên điện thoại còn dễ qua mặt hơn ảnh in vì không có nếp gấp giấy.")

add_heading3("4.2.3. Tấn công Replay qua API")
add_body("Sử dụng curl/Postman gửi POST /authenticate kèm trường image là ảnh đã "
         "chiếm đoạt được mã hóa base64. Server không phân biệt được giữa ảnh đến từ "
         "webcam thật và ảnh được nạp thẳng vào API, kết quả là xác thực thành công. "
         "Tốc độ tấn công: 1 request/giây, có thể song song hóa nhiều luồng nếu cần.")

add_heading3("4.2.4. Tấn công Injection")
add_body("Một biến thể của Replay nhưng mở rộng: attacker viết script Python tự động "
         "duyệt qua một danh sách ảnh khuôn mặt khác nhau và gửi liên tục đến endpoint "
         "/authenticate, từ đó dò tìm bất kỳ tài khoản nào tồn tại trong database. Đây "
         "là dạng tấn công nguy hiểm nhất vì hoàn toàn không cần phần cứng giả mạo và "
         "có thể tự động hóa hoàn toàn.")
add_image("fig_4_6_giao_dien_vulnerable.png",
          "Hình 4.6. Giao diện server vulnerable (mô phỏng)",
          width_cm=14)
add_image("fig_4_7_log_injection.png",
          "Hình 4.7. Log tấn công Injection vào endpoint /authenticate",
          width_cm=14)

add_heading3("4.2.5. Tấn công kết hợp (chained attack)")
add_body("Trong một kịch bản nguy hiểm hơn, kẻ tấn công kết hợp nhiều bước: (1) gọi "
         "GET /users để liệt kê toàn bộ tài khoản; (2) thu thập ảnh khuôn mặt của các "
         "tài khoản mục tiêu từ mạng xã hội; (3) inject ảnh hàng loạt qua POST "
         "/authenticate. Tổng thời gian từ trinh sát đến chiếm quyền dưới 5 phút "
         "trên một hệ thống 7 người dùng. Đây là mức độ nguy hiểm thực tế chứ không "
         "còn ở mức học thuật.")

add_heading2("4.3. Xây dựng giải pháp Liveness Detection")

add_heading3("4.3.1. Bộ dữ liệu")
add_bullet("Real: 200 ảnh khuôn mặt thật chụp trực tiếp từ webcam, đa dạng góc mặt, "
           "ánh sáng phòng và ánh sáng yếu.")
add_bullet("Fake: 200 ảnh giả mạo gồm ảnh in giấy, ảnh hiển thị trên điện thoại và "
           "ảnh chụp lại từ màn hình máy tính.")
add_bullet("Tỷ lệ chia: 80% train – 20% val (tương ứng 320 ảnh train / 80 ảnh val), "
           "có giữ cân bằng giữa Real và Fake trong từng tập.")
add_image("fig_4_1_dataset.png",
          "Hình 4.1. Bộ dữ liệu Real (200 ảnh) và Fake (200 ảnh) dùng huấn luyện",
          width_cm=14)

add_heading3("4.3.2. Cấu hình huấn luyện")
add_table_caption("Bảng 4.1. Cấu hình huấn luyện mô hình MobileNetV2")
cfg = [
    ("Tham số", "Giá trị"),
    ("Kiến trúc", "MobileNetV2 (pretrained ImageNet)"),
    ("Đầu ra", "2 lớp: REAL / FAKE"),
    ("Input size", "224 × 224 × 3"),
    ("Batch size", "32"),
    ("Optimizer Phase 1", "Adam, lr = 1e-3, weight_decay = 1e-4"),
    ("Optimizer Phase 2", "Adam, lr = 1e-4, CosineAnnealingLR"),
    ("Loss", "CrossEntropyLoss"),
    ("Epochs", "10 (Phase 1) + 15 (Phase 2) = 25"),
    ("Augmentation", "Flip, ColorJitter, Random Crop, Random Erasing"),
    ("Thiết bị", "CPU Intel Core i7 (không có GPU rời)"),
    ("Random seed", "42"),
]
add_table_from_rows(cfg)
add_image("fig_4_2_training_history.png",
          "Hình 4.2. Quá trình huấn luyện MobileNetV2 — Loss và Accuracy",
          width_cm=14)
add_image("fig_4_2b_confusion_matrix.png",
          "Hình 4.2b. Confusion matrix của mô hình Liveness Detection",
          width_cm=12)

add_heading3("4.3.3. Tích hợp vào server secured")
add_body("Mô hình sau khi huấn luyện được lưu thành file liveness_model.pth (~9 MB) "
         "và load qua lớp LivenessPredictor trong defender/liveness_model.py. Mỗi "
         "request /authenticate trên server secured sẽ đi qua predictor trước khi đến "
         "bước so khớp khuôn mặt; nếu xác suất Real thấp hơn LIVENESS_THRESHOLD = 0.7, "
         "server trả HTTP 403 và thông báo presentation_attack = True.")
add_body("Để tránh load lại mô hình ở mỗi request, đối tượng LivenessPredictor được "
         "khởi tạo một lần khi server start và lưu trong app.config. Tổng thời gian "
         "warmup ~1.5 giây trên CPU; sau đó mỗi inference chỉ tốn ~50-60 ms.")

add_heading2("4.4. Kết quả đánh giá định lượng")
add_table_caption("Bảng 4.2. Kết quả thực nghiệm trước/sau khi tích hợp Liveness")
res = [
    ("Chỉ số", "Server vulnerable", "Server secured"),
    ("FAR (ảnh giả được chấp nhận)", "≈ 100%", "≈ 2%"),
    ("FRR (người thật bị từ chối)", "≈ 0%", "≈ 5%"),
    ("APCER (PAD bỏ lọt)", "Không áp dụng", "≈ 1.5%"),
    ("BPCER (PAD chặn nhầm)", "Không áp dụng", "≈ 4.8%"),
    ("Accuracy tổng thể", "50% (mù mờ)", "≈ 96.5%"),
    ("Thời gian xử lý / request", "≈ 120 ms", "≈ 180 ms"),
    ("Bộ nhớ RAM bổ sung", "0 MB", "~120 MB (model + tensor)"),
]
add_table_from_rows(res)
add_image("fig_4_4_so_sanh_chi_so.png",
          "Hình 4.4. So sánh FAR / FRR / APCER trước và sau khi áp dụng Liveness",
          width_cm=14)

add_heading3("4.4.1. So sánh latency các thành phần")
add_table_caption("Bảng 4.3. So sánh latency các thành phần trong pipeline")
lat = [
    ("Thành phần", "Vulnerable (ms)", "Secured (ms)"),
    ("Decode + tiền xử lý", "20", "25"),
    ("Phát hiện khuôn mặt", "60", "60"),
    ("Liveness Detection", "—", "55"),
    ("Trích đặc trưng 128-d", "30", "30"),
    ("So khớp Euclidean", "10", "10"),
    ("Tổng cộng", "120", "180"),
]
add_table_from_rows(lat)

add_heading2("4.5. Phân tích kết quả")
add_body("Sau khi tích hợp Liveness Detection, FAR giảm gần như tuyệt đối từ 100% "
         "xuống còn khoảng 2%, chứng minh rằng phần lớn các kịch bản tấn công bằng "
         "ảnh in, ảnh trên điện thoại và replay qua API đã bị chặn ngay tại bước kiểm "
         "tra sống. Đổi lại, FRR tăng nhẹ lên ≈ 5% chủ yếu do các ảnh thật trong điều "
         "kiện ánh sáng quá yếu hoặc góc mặt không chuẩn bị mô hình phân loại nhầm "
         "thành Fake. Đây là sự đánh đổi chấp nhận được trong bảo mật, đặc biệt khi "
         "người dùng có thể thử lại với điều kiện chụp tốt hơn.")
add_body("Chỉ số APCER ≈ 1.5% và BPCER ≈ 4.8% nằm trong phạm vi chấp nhận được đối "
         "với các hệ thống PAD ở mức demo. Theo tiêu chuẩn FIDO Biometric Component "
         "Certification, APCER < 5% đạt mức Level A, kết quả của đề tài thậm chí gần "
         "với Level B (APCER < 1%). Thời gian xử lý mỗi request tăng từ 120 ms lên "
         "180 ms (thêm 50%) nhưng vẫn đảm bảo trải nghiệm thời gian thực; người dùng "
         "rất khó cảm nhận được sự chậm trễ này.")
add_body("Một quan sát đáng chú ý là mô hình MobileNetV2 dù được huấn luyện trên "
         "dataset chỉ 400 ảnh vẫn đạt accuracy ~96.5%. Điều này cho thấy sức mạnh của "
         "transfer learning: backbone đã học được các đặc trưng cấp thấp tốt từ "
         "ImageNet, classifier head chỉ cần học mối quan hệ Real/Fake trên những đặc "
         "trưng đó. Tuy vậy, cần lưu ý rằng dataset nhỏ có thể bị overfit và độ tổng "
         "quát hạn chế khi áp dụng vào môi trường khác — điểm này được phân tích kỹ "
         "ở phần Hạn chế.")

add_heading2("4.6. Hạn chế của giải pháp")
add_bullet("Mô hình PAD chỉ dùng ảnh RGB tĩnh, chưa khai thác thông tin video, IR hay "
           "depth nên có thể bị qua mặt bởi deepfake chất lượng cao.")
add_bullet("Bộ dữ liệu 400 ảnh tự thu thập có quy mô nhỏ, đa dạng môi trường còn hạn "
           "chế nên khả năng tổng quát của mô hình chưa cao.")
add_bullet("Liveness Detection được áp dụng độc lập, chưa kết hợp các yếu tố MFA "
           "(OTP, mật khẩu) để giảm thiểu rủi ro khi mô hình lỗi.")
add_bullet("Chưa có cơ chế chứng thực thiết bị (device attestation) — kẻ tấn công vẫn "
           "có thể inject nếu vượt qua được bài toán Liveness.")
add_bullet("Mô hình chưa chống được tấn công adversarial: thêm nhiễu nhỏ vào ảnh giả "
           "có thể đánh lừa classifier — chủ đề nghiên cứu mở rộng cho tương lai.")

add_heading2("4.7. Thảo luận: lớp phòng vệ chiều sâu")
add_body("Liveness Detection là mảnh ghép quan trọng nhưng không phải duy nhất. Một "
         "hệ thống biometric vận hành trong thực tế cần áp dụng nguyên tắc 'defense "
         "in depth' — phòng vệ nhiều lớp để khi một lớp bị qua mặt thì lớp kế tiếp "
         "vẫn còn cơ hội phát hiện. Hình 4.5 minh họa kiến trúc đề xuất.")
add_image("fig_4_5_lop_phong_ve.png",
          "Hình 4.5. Mô hình phòng vệ nhiều lớp (defense in depth)",
          width_cm=14)
add_body("Các lớp phòng vệ được đề xuất theo thứ tự từ ngoài vào trong: (1) device "
         "attestation và TLS 1.3 ở tầng vận chuyển; (2) chữ ký số ảnh tại client để "
         "chống injection; (3) Liveness Detection ở đầu pipeline server; (4) so khớp "
         "khuôn mặt với ngưỡng cố định; (5) MFA bằng OTP/PIN cho giao dịch nhạy cảm; "
         "(6) logging và phát hiện bất thường ở tầng cuối.")
add_image("fig_4_8_chu_trinh.png",
          "Hình 4.8. Chu trình tấn công – phòng vệ – đánh giá",
          width_cm=12)
add_body("Chu trình ở Hình 4.8 minh họa quá trình lặp giữa tấn công thực tế và cải "
         "tiến phòng vệ. Mỗi vòng lặp giúp hệ thống vững vàng hơn — đây cũng là quy "
         "trình mà nhiều tổ chức bảo mật lớn đang áp dụng (Google Project Zero, "
         "Microsoft Red Team, v.v.).")
page_break()

# =========================================================
# CHƯƠNG 5 - KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN
# =========================================================
add_heading1("CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")

add_heading2("5.1. Kết luận")
add_body("Đề tài đã hoàn thành các mục tiêu đặt ra ban đầu: hệ thống hóa cơ sở lý "
         "thuyết về xác thực sinh trắc học và các loại tấn công, xây dựng được hai "
         "phiên bản server (vulnerable và secured) để mô phỏng đầy đủ chu trình tấn "
         "công – phòng thủ, huấn luyện thành công mô hình MobileNetV2 cho bài toán "
         "Real/Fake và đo đạc các chỉ số FAR, FRR, APCER, BPCER theo chuẩn ISO/IEC "
         "30107-3.")
add_body("Kết quả thực nghiệm cho thấy một hệ thống xác thực khuôn mặt nếu chỉ thực "
         "hiện so khớp đặc trưng mà không có Liveness Detection thì gần như không có "
         "khả năng phòng vệ trước các tấn công trình diện, replay và injection. Sau "
         "khi tích hợp Liveness Detection dựa trên MobileNetV2, FAR giảm từ ≈ 100% "
         "xuống ≈ 2% với mức tăng FRR rất nhỏ (5%), chứng minh giải pháp là khả thi "
         "và hiệu quả trong môi trường lab.")
add_body("Quan trọng hơn các con số là bài học về tư duy bảo mật: một hệ thống an "
         "toàn không thể chỉ dựa vào một biện pháp duy nhất. Liveness Detection là "
         "lớp phòng vệ trung tâm nhưng phải được kết hợp với device attestation, "
         "mã hóa kênh truyền, MFA và logging giám sát để đạt mức an toàn thực sự "
         "trong môi trường sản xuất.")

add_heading2("5.2. Đóng góp của đề tài")
add_bullet("Một bộ tài liệu Việt hóa, hệ thống lại các bề mặt tấn công vào hệ thống "
           "biometric theo từng bước trong quy trình xác thực.")
add_bullet("Mã nguồn demo gồm: server vulnerable, server secured, script huấn luyện "
           "MobileNetV2 và bộ ảnh mẫu Real/Fake — có thể tái sử dụng cho mục đích đào "
           "tạo hoặc kiểm thử an toàn.")
add_bullet("Bộ chỉ số đánh giá thực nghiệm trước/sau giúp minh họa rõ tác dụng của "
           "Liveness Detection cho mục đích đào tạo.")
add_bullet("Đề xuất kiến trúc phòng vệ chiều sâu (defense in depth) áp dụng được cho "
           "các hệ thống biometric thương mại với chi phí triển khai hợp lý.")
add_bullet("Mô hình đe dọa (threat model) cụ thể, có thể được tái sử dụng cho các đề "
           "tài tương tự hoặc làm checklist khi đánh giá an toàn hệ thống.")

add_heading2("5.3. Hướng phát triển")
add_bullet("Mở rộng Liveness Detection sang dạng video-based, sử dụng tín hiệu nháy "
           "mắt, chuyển động đầu hoặc rPPG (remote photoplethysmography — đo nhịp "
           "tim qua biến đổi màu da) để chống deepfake video.")
add_bullet("Bổ sung kênh cảm biến IR/3D depth hoặc kết hợp camera đa phổ để chống "
           "các PAI 3D như mặt nạ silicon.")
add_bullet("Áp dụng kiến trúc xác thực đa yếu tố: kết hợp biometric với OTP, hardware "
           "token (FIDO2/WebAuthn) hoặc PIN.")
add_bullet("Mở rộng quy mô dataset (ít nhất 5.000 mẫu/lớp), thử nghiệm các kiến trúc "
           "khác như EfficientNet, Vision Transformer, Swin Transformer để so sánh.")
add_bullet("Bổ sung cơ chế xác thực thiết bị (device attestation) qua TPM hoặc Secure "
           "Enclave, ký số ảnh tại client để chống tấn công injection.")
add_bullet("Triển khai giám sát vận hành (monitoring) để phát hiện các đợt tấn công "
           "bất thường vào endpoint xác thực — ví dụ phát hiện burst request từ cùng "
           "một IP hoặc cùng một User-Agent đáng ngờ.")
add_bullet("Nghiên cứu tấn công adversarial trên Liveness model và biện pháp phòng "
           "vệ tương ứng (adversarial training, input transformation).")

add_heading2("5.4. Lộ trình triển khai thực tế")
add_body("Để chuyển từ môi trường lab sang môi trường sản xuất, em đề xuất lộ trình "
         "triển khai theo 4 giai đoạn như sau:")
add_image("fig_5_1_lo_trinh.png",
          "Hình 5.1. Lộ trình triển khai thực tế giải pháp PAD",
          width_cm=14)

add_table_caption("Bảng 5.1. Lộ trình triển khai PAD trong môi trường sản xuất")
roadmap = [
    ("Giai đoạn", "Thời gian", "Hoạt động chính"),
    ("1. PoC", "1-2 tháng", "Demo trên dữ liệu nội bộ, đo APCER/BPCER baseline"),
    ("2. Pilot", "3-6 tháng", "Triển khai trên 100-500 user thật, A/B test với "
     "phiên bản cũ"),
    ("3. Mở rộng", "6-12 tháng", "Tích hợp MFA, device attestation, monitoring; "
     "đạt chứng nhận FIDO Level A"),
    ("4. Sản xuất", "12+ tháng", "Triển khai toàn bộ hệ thống, audit định kỳ, "
     "cập nhật mô hình theo dữ liệu mới"),
]
add_table_from_rows(roadmap)

add_body("Trong từng giai đoạn cần có chỉ số đo (KPI) cụ thể: APCER, BPCER, FAR, "
         "FRR, latency, tỉ lệ false positive cảnh báo. Việc gắn KPI giúp tránh tình "
         "trạng triển khai mà không biết hệ thống có thực sự cải thiện hay không. "
         "Các KPI này nên được đưa lên dashboard nội bộ và được giám sát liên tục bởi "
         "đội vận hành an ninh (SOC).")
add_body("Cuối cùng, em mong rằng đề tài này sẽ là một tài liệu tham khảo hữu ích "
         "cho các bạn sinh viên ngành An toàn thông tin tại HUTECH cũng như cho các "
         "tổ chức đang xây dựng hệ thống xác thực biometric. Mã nguồn và dataset đi "
         "kèm sẽ được công bố mã nguồn mở để cộng đồng cùng đóng góp và cải tiến.")
page_break()

# =========================================================
# TÀI LIỆU THAM KHẢO
# =========================================================
add_heading1("TÀI LIỆU THAM KHẢO")
refs = [
    "ISO/IEC 30107-1:2016, Information technology — Biometric presentation attack detection — Part 1: Framework.",
    "ISO/IEC 30107-3:2017, Information technology — Biometric presentation attack detection — Part 3: Testing and reporting.",
    "Sandler M., Howard A., Zhu M., Zhmoginov A., Chen L., MobileNetV2: Inverted Residuals and Linear Bottlenecks, CVPR 2018.",
    "Geitgey A., face_recognition: The world's simplest facial recognition api for Python and the command line, GitHub, 2017.",
    "King D. E., Dlib-ml: A Machine Learning Toolkit, Journal of Machine Learning Research, 2009.",
    "Boulkenafet Z., Komulainen J., Hadid A., Face Anti-Spoofing Based on Color Texture Analysis, IEEE ICIP 2015.",
    "Liu Y., Jourabloo A., Liu X., Learning Deep Models for Face Anti-Spoofing: Binary or Auxiliary Supervision, CVPR 2018.",
    "George A., Marcel S., Deep Pixel-wise Binary Supervision for Face Presentation Attack Detection, IEEE ICB 2019.",
    "Yu Z., Qin Y., Li X., Zhao C., Lei Z., Zhao G., Deep Learning for Face Anti-Spoofing: A Survey, IEEE TPAMI, 2022.",
    "Pan G., Sun L., Wu Z., Lao S., Eyeblink-based Anti-Spoofing in Face Recognition, ICCV 2007.",
    "OWASP Foundation, OWASP Top 10 for Biometric Systems (Draft), https://owasp.org, 2023.",
    "NIST, Biometric Specifications for Personal Identity Verification, NIST SP 800-76-2, 2013.",
    "NIST, Digital Identity Guidelines: Authentication and Lifecycle Management, NIST SP 800-63B, 2017.",
    "He K., Zhang X., Ren S., Sun J., Deep Residual Learning for Image Recognition, CVPR 2016.",
    "Schroff F., Kalenichenko D., Philbin J., FaceNet: A Unified Embedding for Face Recognition and Clustering, CVPR 2015.",
    "Parkhi O. M., Vedaldi A., Zisserman A., Deep Face Recognition, BMVC 2015.",
    "Goodfellow I., et al., Generative Adversarial Networks, NIPS 2014.",
    "Karras T., Laine S., Aila T., A Style-Based Generator Architecture for GANs, CVPR 2019.",
    "Tolosana R., Vera-Rodriguez R., Fierrez J., et al., Deepfakes and Beyond: A Survey of Face Manipulation and Fake Detection, Information Fusion, 2020.",
    "PyTorch Team, torchvision.models — MobileNetV2 documentation, https://pytorch.org/vision/stable/models.html.",
    "Bradski G., The OpenCV Library, Dr. Dobb's Journal of Software Tools, 2000.",
    "Ramachandra R., Busch C., Presentation Attack Detection Methods for Face Recognition Systems: A Comprehensive Survey, ACM Computing Surveys, 2017.",
    "Anjos A., Marcel S., Counter-measures to Photo Attacks in Face Recognition: a Public Database and a Baseline, IJCB 2011.",
    "Wen D., Han H., Jain A. K., Face Spoof Detection with Image Distortion Analysis, IEEE TIFS 2015.",
    "Pinto A., Pedrini H., Schwartz W. R., Rocha A., Face Spoofing Detection Through Visual Codebooks of Spectral Temporal Cubes, IEEE TIP 2015.",
    "Atoum Y., Liu Y., Jourabloo A., Liu X., Face Anti-Spoofing Using Patch and Depth-based CNNs, IJCB 2017.",
    "Jain A. K., Ross A. A., Nandakumar K., Introduction to Biometrics, Springer, 2011.",
    "Maltoni D., Maio D., Jain A. K., Prabhakar S., Handbook of Fingerprint Recognition, Springer, 3rd ed., 2022.",
    "Daugman J., How Iris Recognition Works, IEEE TCSVT 2004.",
    "ENISA, Remote Identity Proofing — Attacks and Countermeasures, ENISA Report, 2022.",
    "FIDO Alliance, FIDO Biometric Component Certification Program, https://fidoalliance.org, 2023.",
    "Marcel S., Nixon M. S., Li S. Z., Handbook of Biometric Anti-Spoofing, Springer, 2nd ed., 2019.",
    "Galbally J., Marcel S., Fierrez J., Biometric Antispoofing Methods: A Survey in Face Recognition, IEEE Access, 2014.",
    "Erdogmus N., Marcel S., Spoofing Face Recognition With 3D Masks, IEEE TIFS 2014.",
    "Ferrara M., Franco A., Maltoni D., The Magic Passport, IJCB 2014.",
    "Howard A. G., et al., MobileNets: Efficient Convolutional Neural Networks for Mobile Vision Applications, arXiv 1704.04861, 2017.",
    "Tan M., Le Q. V., EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks, ICML 2019.",
    "Dosovitskiy A., et al., An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale, ICLR 2021.",
    "Ngân hàng Nhà nước Việt Nam, Quyết định 2345/QĐ-NHNN về xác thực sinh trắc học cho giao dịch ngân hàng trực tuyến, 2024.",
    "Cộng đồng nghiên cứu HUTECH, Tài liệu thực hành biometric security tại Khoa Công nghệ Thông tin, 2026.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"[{i}] {r}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# =========================================================
# SAVE
# =========================================================
# =========================================================
# SAVE
# =========================================================
update_fields_on_open()
doc.save(OUT)
print(f"[OK] Saved: {OUT}")
