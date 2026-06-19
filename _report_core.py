# -*- coding: utf-8 -*-
"""_report_core.py - helpers cho build_final_report_v2.py.

Sua doi v3:
- Tat ca chu va bang deu mau den.
- Heading dung Word's built-in style => Ctrl-F + Navigation Pane hoat dong.
- TOC dung w:fldChar SDT thay vi static text => co the auto-update F9.
- Code block boc trong 1-cell table monospace co border + light grey background.
- Figures/Tables index: track page so xap xi.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_BREAK,
)
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

ASSETS = r"C:\Users\ADMIN\biometric-security-lab\_report_assets"
OUT = r"C:\Users\ADMIN\Downloads\BAOCAODOANCOSO_NGUYENPHUCTHINH_FINAL_v2.docx"

FIGURES_LOG = []   # list of (num, caption)
TABLES_LOG = []    # list of (num, title)
PAGE_TRACKER = {"counter": 0}  # so paragraph dem -> uoc tinh trang


# ============================================================
# Run / paragraph
# ============================================================
def set_run_font(run, size=13, bold=False, italic=False,
                 font='Times New Roman', color='000000'):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc_or_cell, text, *, size=13, bold=False, italic=False,
             align='justify', indent_first=0.0, space_after=6,
             line_spacing=1.5, color='000000'):
    p = doc_or_cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    am = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
          'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
    p.alignment = am.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    # estimate page contribution
    try:
        n_lines = max(1, len(text) // 80)
    except Exception:
        n_lines = 1
    PAGE_TRACKER["counter"] += n_lines
    return p


# ============================================================
# Headings - dung Word built-in style de Navigation Pane / TOC nhin thay
# ============================================================
def _heading(doc, text, *, level, size, align='left', upper=False, italic=False):
    style_name = f'Heading {level}'
    p = doc.add_paragraph(style=style_name)
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(14)
        pf.space_after = Pt(10)
        PAGE_TRACKER["counter"] += 20  # h1 = phan dau chuong, gan +0.5 trang
    elif level == 2:
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
        PAGE_TRACKER["counter"] += 2
    else:
        pf.space_before = Pt(6)
        pf.space_after = Pt(4)
        PAGE_TRACKER["counter"] += 1
    pf.line_spacing = 1.5
    align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                 'center': WD_ALIGN_PARAGRAPH.CENTER,
                 'right': WD_ALIGN_PARAGRAPH.RIGHT}
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    txt = text.upper() if upper else text
    run = p.add_run(txt)
    set_run_font(run, size=size, bold=True, italic=italic, color='000000')
    return p


def add_h1(doc, text):
    return _heading(doc, text, level=1, size=15, align='center', upper=True)


def add_h2(doc, text):
    return _heading(doc, text, level=2, size=13, align='left')


def add_h3(doc, text):
    return _heading(doc, text, level=3, size=13, align='left', italic=True)


def add_bullet(doc, text, *, size=13):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=size, color='000000')
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color='000000')
    PAGE_TRACKER["counter"] += max(1, len(text) // 80)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ============================================================
# Tables
# ============================================================
def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, color='000000', sz='6'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:color'), color)
        tcBorders.append(e)
    tcPr.append(tcBorders)


def add_table(doc, headers, rows, widths=None, header_color='D9D9D9'):
    """Bang nen header xam, chu DEN, vien den."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    # header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        para = c.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h)
        set_run_font(r, size=12, bold=True, color='000000')
        shade_cell(c, header_color)
        set_cell_border(c)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = ''
            para = c.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.2
            r = para.add_run(str(val))
            set_run_font(r, size=11, color='000000')
            set_cell_border(c)
            if ri % 2 == 0:
                shade_cell(c, 'F5F5F5')
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


# ============================================================
# Code block - bang 1-cell, monospace, light grey background, vien den
# ============================================================
def add_code_block(doc, code_text, *, size=10):
    """Boc code trong khung mau xam nhat."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    cell = t.rows[0].cells[0]
    cell.text = ''
    cell.width = Cm(15.5)
    shade_cell(cell, 'F2F2F2')
    set_cell_border(cell, color='808080', sz='8')
    for line in code_text.split('\n'):
        para = cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(0)
        r = para.add_run(line if line else ' ')
        set_run_font(r, size=size, font='Consolas', color='000000')
    # remove trang dau cua cell (do .text='' tao 1 paragraph trong)
    first = cell.paragraphs[0]
    if not first.text:
        first._element.getparent().remove(first._element)
    return t


# ============================================================
# Image / table caption
# ============================================================
def _approx_page():
    # Uoc tinh trang dua tren PAGE_TRACKER, ratio ~ 35 paragraph / page
    return max(1, PAGE_TRACKER["counter"] // 35 + 1)


def _bump_para_counter(n=1):
    PAGE_TRACKER["counter"] += n


def add_image(doc, filename, width_cm=14, fig_num=None, caption=None):
    path = os.path.join(ASSETS, filename) if not os.path.isabs(filename) else filename
    if not os.path.exists(path):
        add_para(doc, f"[Hình {fig_num}: {caption}]",
                 align='center', italic=True, color='666666')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption and fig_num:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run(f"Hình {fig_num}. {caption}")
        set_run_font(r, size=11, italic=True, bold=True, color='000000')
        FIGURES_LOG.append((fig_num, caption, _approx_page()))
    _bump_para_counter(8)  # 1 image ~ 0.25 page


def add_table_caption(doc, num, title):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(f"Bảng {num}. {title}")
    set_run_font(r, size=11, italic=True, bold=True, color='000000')
    TABLES_LOG.append((num, title, _approx_page()))
    _bump_para_counter(4)


# ============================================================
# Section helpers - Roman / Decimal page numbering
# ============================================================
def _add_pgNumType(section, fmt, start=1):
    sectPr = section._sectPr
    existing = sectPr.find(qn('w:pgNumType'))
    if existing is not None:
        sectPr.remove(existing)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)


def _add_page_field_to_footer(section, prefix=""):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        if p.text.strip() == "":
            continue
        p._element.getparent().remove(p._element)
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if prefix:
        run_pre = p.add_run(prefix)
        set_run_font(run_pre, size=12, color='000000')
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE \\* MERGEFORMAT')
    r_in = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    r_in.append(rPr)
    t = OxmlElement('w:t')
    t.text = '1'
    r_in.append(t)
    fldSimple.append(r_in)
    p._p.append(fldSimple)


def setup_doc_defaults(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style.font.color.rgb = RGBColor(0, 0, 0)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), 'Times New Roman')

    # Force Heading 1/2/3 to BLACK Times New Roman
    for lvl, sz in [(1, 15), (2, 13), (3, 13)]:
        try:
            h = doc.styles[f'Heading {lvl}']
            h.font.name = 'Times New Roman'
            h.font.size = Pt(sz)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0, 0, 0)
            rPr2 = h.element.get_or_add_rPr()
            rFonts2 = rPr2.find(qn('w:rFonts'))
            if rFonts2 is None:
                rFonts2 = OxmlElement('w:rFonts')
                rPr2.append(rFonts2)
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                rFonts2.set(qn(attr), 'Times New Roman')
        except KeyError:
            pass


def new_section(doc, fmt='upperRoman', start=1):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.2)
    _add_pgNumType(section, fmt, start)
    _add_page_field_to_footer(section)
    return section


# ============================================================
# TOC field - real Word TOC, F9 to update
# ============================================================
def add_toc_field(doc, levels='1-3'):
    """Insert real TOC field that Word can update with F9."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f'TOC \\o "{levels}" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = ('Right-click here and choose "Update Field" '
                     '(or press F9) to populate the Table of Contents.')

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._element
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)

    # Tell Word to update fields on open
    settings = doc.settings.element
    update = settings.find(qn('w:updateFields'))
    if update is None:
        update = OxmlElement('w:updateFields')
        update.set(qn('w:val'), 'true')
        settings.append(update)

    return p
