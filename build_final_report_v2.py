# -*- coding: utf-8 -*-
"""build_final_report_v2.py
Orchestrator chinh - sinh BAOCAODOANCOSO_NGUYENPHUCTHINH_FINAL_v2.docx.

- Section 1 (cover): khong co so trang
- Section 2 (front matter): so trang La Ma I, II, III...
- Section 3 (chuong + tai lieu tham khao): so trang 1, 2, 3...

Doc Danh muc Hinh / Bang duoc dien tu dong bang two-pass build:
  Pass 1 - chay du chuong vao doc tam de populate FIGURES_LOG/TABLES_LOG
  Pass 2 - reset log, build doc that voi Danh muc da co san
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

import _report_core
from _report_core import (
    setup_doc_defaults, new_section, OUT,
    add_para, add_h1, add_h2, add_h3, add_bullet, add_table,
    add_page_break, set_run_font, add_toc_field,
)
from _report_chapter1 import build_chapter1
from _report_chapter2 import build_chapter2
from _report_chapter3 import build_chapter3
from _report_chapter4 import build_chapter4
from _report_chapter5 import build_chapter5, build_references


# ============================================================
# COVER
# ============================================================
def build_cover(doc):
    add_para(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO",
             size=13, bold=True, align='center', space_after=2)
    add_para(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HCM (HUTECH)",
             size=13, bold=True, align='center', space_after=2)
    add_para(doc, "KHOA CÔNG NGHỆ THÔNG TIN",
             size=13, bold=True, align='center', space_after=24)

    add_para(doc, "ĐỒ ÁN CƠ SỞ",
             size=24, bold=True, align='center', space_after=14)
    add_para(doc,
        "NGHIÊN CỨU KỸ THUẬT PHÁT HIỆN\nTẤN CÔNG TRÌNH DIỆN TRONG HỆ THỐNG\nXÁC THỰC SINH TRẮC HỌC",
        size=18, bold=True, align='center', space_after=24, line_spacing=1.3)

    add_para(doc, "Ngành: AN TOÀN THÔNG TIN",
             size=14, bold=True, align='center', space_after=22)

    add_para(doc, "Giảng viên hướng dẫn: ThS. Đặng Thị Thạch Thảo",
             size=13, align='center', space_after=4)
    add_para(doc, "Sinh viên thực hiện: Nguyễn Phúc Thịnh",
             size=13, align='center', space_after=4)
    add_para(doc, "MSSV: 2387700066    Lớp: 23DATA1",
             size=13, align='center', space_after=40)

    add_para(doc, "TP. Hồ Chí Minh, tháng 6 năm 2026",
             size=13, italic=True, align='center')


# ============================================================
# FRONT MATTER
# ============================================================
def build_loi_mo_dau(doc):
    add_h1(doc, "LỜI MỞ ĐẦU")
    add_para(doc,
        "Trong hơn một thập kỷ trở lại đây, xác thực sinh trắc học (biometric "
        "authentication) đã trở thành phương thức xác thực chủ đạo, được tích hợp rộng "
        "rãi trên điện thoại thông minh, hệ thống ngân hàng điện tử, kiểm soát ra vào "
        "và các dịch vụ định danh công dân điện tử. So với mật khẩu truyền thống, sinh "
        "trắc học mang lại trải nghiệm thuận tiện hơn — người dùng không cần ghi nhớ "
        "thông tin và tốc độ xác thực chỉ trong vài trăm mili-giây [1][12].",
        indent_first=1.0)
    add_para(doc,
        "Tuy nhiên, đi cùng với sự phổ biến đó là các nguy cơ bảo mật ngày càng nghiêm "
        "trọng. Các hệ thống nhận diện khuôn mặt nếu không được trang bị cơ chế kiểm "
        "tra sống (Liveness Detection / Presentation Attack Detection) rất dễ bị qua "
        "mặt bằng ảnh in, ảnh hiển thị trên điện thoại, video phát lại hoặc các kỹ "
        "thuật injection trực tiếp ở tầng API. Khác với mật khẩu, dữ liệu sinh trắc "
        "học một khi bị lộ thì không thể thay đổi, vì vậy hậu quả của một cuộc tấn "
        "công thành công là rất lâu dài [1][9][29].",
        indent_first=1.0)
    add_para(doc,
        "Đề tài “Nghiên cứu kỹ thuật phát hiện tấn công trình diện trong hệ thống xác "
        "thực sinh trắc học” được thực hiện nhằm khảo sát các bề mặt tấn công của một "
        "hệ thống xác thực khuôn mặt thực tế dưới góc nhìn của một SOC Analyst — coi "
        "mỗi frame ảnh là một telemetry record và áp dụng vòng đời PAD Engineering để "
        "thiết kế lớp Detection Engine. Báo cáo trình bày toàn bộ quá trình từ cơ sở "
        "lý thuyết, Threat Modeling, Red Teaming, đến huấn luyện Detection Engine "
        "MobileNetV2 và đánh giá định lượng theo tiêu chuẩn ISO/IEC 30107-3.",
        indent_first=1.0)
    add_page_break(doc)


def build_loi_cam_doan(doc):
    add_h1(doc, "LỜI CAM ĐOAN")
    add_para(doc,
        "Em xin cam đoan đồ án “Nghiên cứu kỹ thuật phát hiện tấn công trình diện "
        "trong hệ thống xác thực sinh trắc học” là công trình nghiên cứu của riêng "
        "em, được thực hiện dưới sự hướng dẫn của giảng viên ThS. Đặng Thị Thạch "
        "Thảo. Các số liệu, kết quả thực nghiệm trong báo cáo là trung thực, được đo "
        "đạc trên môi trường lab do em tự xây dựng và chưa từng được công bố trong "
        "bất kỳ công trình nào khác.",
        indent_first=1.0)
    add_para(doc,
        "Các nguồn tài liệu tham khảo đều được trích dẫn rõ ràng theo định dạng [n]. "
        "Toàn bộ thí nghiệm Red Teaming chỉ được thực hiện trên hệ thống do em tự "
        "triển khai phục vụ mục đích nghiên cứu, không nhằm vào bất kỳ hệ thống "
        "thương mại của tổ chức hoặc cá nhân khác.",
        indent_first=1.0)
    add_para(doc, "TP. HCM, ngày … tháng … năm 2026",
             align='right', italic=True, space_after=2)
    add_para(doc, "Sinh viên thực hiện",
             align='right', italic=True, space_after=20)
    add_para(doc, "Nguyễn Phúc Thịnh", align='right', bold=True)
    add_page_break(doc)


def build_loi_cam_on(doc):
    add_h1(doc, "LỜI CẢM ƠN")
    add_para(doc,
        "Trước tiên, em xin gửi lời cảm ơn chân thành đến quý thầy cô Khoa Công nghệ "
        "Thông tin – Trường Đại học Công nghệ TP. HCM (HUTECH) đã trang bị cho em "
        "những kiến thức nền tảng quan trọng về an toàn thông tin, học máy và phát "
        "triển ứng dụng trong suốt thời gian học tập tại trường.",
        indent_first=1.0)
    add_para(doc,
        "Em xin gửi lời cảm ơn sâu sắc nhất đến giảng viên hướng dẫn ThS. Đặng Thị "
        "Thạch Thảo đã tận tình góp ý, định hướng và tạo điều kiện để em có thể hoàn "
        "thành đề tài. Những phản hồi kịp thời của Cô trong từng tuần báo cáo đã "
        "giúp em điều chỉnh hướng nghiên cứu phù hợp và bám sát mục tiêu đặt ra ban "
        "đầu — đặc biệt là việc tiếp cận bài toán theo tư duy Detection Engineering "
        "thay vì chỉ dừng ở mức nhận diện khuôn mặt thuần túy.",
        indent_first=1.0)
    add_para(doc,
        "Em cũng xin cảm ơn gia đình, bạn bè đã luôn động viên, hỗ trợ em trong quá "
        "trình nghiên cứu và thử nghiệm các kịch bản tấn công, phòng vệ trên hệ "
        "thống sinh trắc học.",
        indent_first=1.0)
    add_para(doc,
        "Do thời gian và kiến thức còn hạn chế, đồ án không tránh khỏi những thiếu "
        "sót. Em rất mong nhận được những góp ý quý báu từ quý thầy cô để hoàn thiện "
        "hơn trong các nghiên cứu tiếp theo. Em xin chân thành cảm ơn!",
        indent_first=1.0)
    add_page_break(doc)


def build_nhan_xet(doc):
    add_h1(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN")
    for _ in range(12):
        add_para(doc, "………………………………………………………………………………………………",
                 align='left', space_after=6, line_spacing=1.8)
    add_para(doc, "TP. HCM, ngày … tháng … năm 2026",
             align='right', italic=True, space_after=4)
    add_para(doc, "Giảng viên hướng dẫn",
             align='right', italic=True, space_after=20)
    add_para(doc, "ThS. Đặng Thị Thạch Thảo",
             align='right', bold=True)
    add_page_break(doc)


def _toc_entry(doc, text, page, level=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.35
    if level == 0:
        pf.left_indent = Cm(0)
    elif level == 1:
        pf.left_indent = Cm(0.6)
    else:
        pf.left_indent = Cm(1.2)
    r = p.add_run(text)
    set_run_font(r, size=12, bold=(level == 0))
    tab = p.add_run("\t")
    set_run_font(tab, size=12)
    pr = p.add_run(str(page))
    set_run_font(pr, size=12, bold=(level == 0))
    pf.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def build_muc_luc(doc):
    add_h1(doc, "MỤC LỤC")
    add_para(doc,
        "Mục lục bên dưới là TOC field tự động — nhấn F9 (hoặc chuột phải → "
        "Update Field) khi mở file để Word liệt kê đầy đủ các Heading và số "
        "trang thực tế. Đồng thời, nhấn Ctrl + click vào tên mục để nhảy nhanh "
        "tới phần tương ứng.",
        size=11, italic=True, indent_first=0, space_after=8, color='555555')
    add_toc_field(doc, levels='1-3')
    add_page_break(doc)


def build_abbreviations(doc):
    add_h1(doc, "DANH MỤC CHỮ VIẾT TẮT")
    rows = [
        ["1",  "APCER",        "Attack Presentation Classification Error Rate"],
        ["2",  "BPCER",        "Bona-fide Presentation Classification Error Rate"],
        ["3",  "FAR",          "False Acceptance Rate"],
        ["4",  "FRR",          "False Rejection Rate"],
        ["5",  "EER",          "Equal Error Rate"],
        ["6",  "PAD",          "Presentation Attack Detection"],
        ["7",  "PAI",          "Presentation Attack Instrument"],
        ["8",  "TTP",          "Tactics, Techniques and Procedures"],
        ["9",  "ASR",          "Attack Success Rate"],
        ["10", "MITRE ATT&CK", "Adversarial Tactics, Techniques, and Common Knowledge"],
        ["11", "STRIDE",       "Spoofing, Tampering, Repudiation, Info Disclosure, DoS, Elevation"],
        ["12", "CNN",          "Convolutional Neural Network"],
        ["13", "MobileNetV2",  "Kiến trúc CNN nhẹ với Inverted Residual + Linear Bottleneck"],
        ["14", "ResNet",       "Residual Network"],
        ["15", "GAN",          "Generative Adversarial Network"],
        ["16", "ROC",          "Receiver Operating Characteristic"],
        ["17", "GAP",          "Global Average Pooling"],
        ["18", "BCE",          "Binary Cross Entropy"],
        ["19", "CLAHE",        "Contrast Limited Adaptive Histogram Equalization"],
        ["20", "ROI",          "Region of Interest"],
        ["21", "RGB",          "Red – Green – Blue"],
        ["22", "ISO/IEC 30107","Tiêu chuẩn quốc tế về Phát hiện tấn công trình diện"],
        ["23", "FIDO",         "Fast IDentity Online Alliance"],
        ["24", "AAL",          "Authenticator Assurance Level (NIST SP 800-63B)"],
        ["25", "OWASP",        "Open Worldwide Application Security Project"],
        ["26", "NIST",         "National Institute of Standards and Technology"],
        ["27", "API",          "Application Programming Interface"],
        ["28", "REST",         "Representational State Transfer"],
        ["29", "JSON",         "JavaScript Object Notation"],
        ["30", "JWT",          "JSON Web Token"],
        ["31", "HMAC",         "Hash-based Message Authentication Code"],
        ["32", "TLS",          "Transport Layer Security"],
        ["33", "MITM",         "Man-in-the-Middle"],
        ["34", "MFA",          "Multi-Factor Authentication"],
        ["35", "SOC",          "Security Operations Center"],
        ["36", "SIEM",         "Security Information and Event Management"],
        ["37", "SOAR",         "Security Orchestration, Automation, and Response"],
        ["38", "HSM",          "Hardware Security Module"],
        ["39", "TEE",          "Trusted Execution Environment"],
        ["40", "DB",           "Database"],
    ]
    add_table(doc, ["STT", "Viết tắt", "Nghĩa đầy đủ"], rows,
              widths=[1.5, 3.5, 11])
    add_page_break(doc)


def build_list_of_figures(doc, figures):
    add_h1(doc, "DANH MỤC HÌNH ẢNH")
    if not figures:
        figures = [("—", "(sẽ được cập nhật sau khi build chương)", "—")]
    rows = []
    for item in figures:
        if len(item) == 3:
            num, title, page = item
        else:
            num, title = item
            page = "—"
        rows.append([f"Hình {num}. {title}", str(page)])
    add_table(doc, ["Tên hình", "Trang"], rows, widths=[14, 2])
    add_page_break(doc)


def build_list_of_tables(doc, tables):
    add_h1(doc, "DANH MỤC BẢNG BIỂU")
    if not tables:
        tables = [("—", "(sẽ được cập nhật sau khi build chương)", "—")]
    rows = []
    for item in tables:
        if len(item) == 3:
            num, title, page = item
        else:
            num, title = item
            page = "—"
        rows.append([f"Bảng {num}. {title}", str(page)])
    add_table(doc, ["Tên bảng", "Trang"], rows, widths=[14, 2])
    add_page_break(doc)


# ============================================================
# DRY-RUN PASS - chi de populate FIGURES_LOG / TABLES_LOG
# ============================================================
def _dry_run_collect():
    dry = Document()
    setup_doc_defaults(dry)
    build_chapter1(dry)
    build_chapter2(dry)
    build_chapter3(dry)
    build_chapter4(dry)
    build_chapter5(dry)
    figures = list(_report_core.FIGURES_LOG)
    tables = list(_report_core.TABLES_LOG)
    _report_core.FIGURES_LOG.clear()
    _report_core.TABLES_LOG.clear()
    return figures, tables


# ============================================================
# MAIN
# ============================================================
def main():
    print("[1/3] Dry-run: thu thap Figures/Tables...")
    figures, tables = _dry_run_collect()
    print(f"     -> {len(figures)} hinh, {len(tables)} bang.")

    print("[2/3] Build doc that...")
    doc = Document()
    setup_doc_defaults(doc)

    # ----- Section 1: Cover (no page numbering) -----
    build_cover(doc)

    # ----- Section 2: Front matter (upperRoman) -----
    new_section(doc, fmt='upperRoman', start=1)
    build_loi_mo_dau(doc)
    build_loi_cam_doan(doc)
    build_loi_cam_on(doc)
    build_nhan_xet(doc)
    build_muc_luc(doc)
    build_abbreviations(doc)
    build_list_of_figures(doc, figures)
    build_list_of_tables(doc, tables)

    # ----- Section 3: Chapters (decimal) -----
    new_section(doc, fmt='decimal', start=1)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_chapter4(doc)
    build_chapter5(doc)
    build_references(doc)

    print("[3/3] Save -> %s" % OUT)
    out_dir = os.path.dirname(OUT)
    if out_dir and not os.path.isdir(out_dir):
        os.makedirs(out_dir, exist_ok=True)
    doc.save(OUT)
    size_kb = os.path.getsize(OUT) / 1024
    print(f"     -> OK, {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
