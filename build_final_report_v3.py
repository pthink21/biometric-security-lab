# -*- coding: utf-8 -*-
"""build_final_report_v3.py - Sinh BAOCAODOANCOSO_HOANCHINH.docx (~70 trang).

Khac biet so voi v2:
- Su dung _v3_chapter1..5 (khong em-dash, khong en-dash trong noi dung).
- Title cover: "NGHIEN CUU TAN CONG VAO HE THONG XAC THUC SINH TRAC HOC".
- Loi mo dau, cam doan, cam on viet lai, sach em-dash.
- Output: BAOCAODOANCOSO_HOANCHINH.docx (cung thu muc project).
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

import _report_core
from _report_core import (
    setup_doc_defaults, new_section,
    add_para, add_h1, add_h2, add_h3, add_bullet, add_table,
    add_page_break, set_run_font, add_toc_field,
)
from _v3_chapter1 import build_chapter1
from _v3_chapter2 import build_chapter2
from _v3_chapter3 import build_chapter3
from _v3_chapter4 import build_chapter4
from _v3_chapter5 import build_chapter5, build_references


OUT_FINAL = r"C:\Users\ADMIN\biometric-security-lab\BAOCAODOANCOSO_HOANCHINH.docx"


# ============================================================
# COVER
# ============================================================
def build_cover(doc):
    add_para(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO",
             size=13, bold=True, align='center', space_after=2)
    add_para(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HCM (HUTECH)",
             size=13, bold=True, align='center', space_after=2)
    add_para(doc, "KHOA CÔNG NGHỆ THÔNG TIN",
             size=13, bold=True, align='center', space_after=24)

    add_para(doc, "ĐỒ ÁN CƠ SỞ",
             size=24, bold=True, align='center', space_after=14)
    add_para(doc,
        "NGHIÊN CỨU TẤN CÔNG VÀO HỆ THỐNG\nXÁC THỰC SINH TRẮC HỌC",
        size=20, bold=True, align='center', space_after=24, line_spacing=1.3)

    add_para(doc, "Ngành: AN TOÀN THÔNG TIN",
             size=14, bold=True, align='center', space_after=22)

    add_para(doc, "Giảng viên hướng dẫn: ThS. Đặng Thị Thạch Thảo",
             size=13, align='center', space_after=4)
    add_para(doc, "Sinh viên thực hiện: Nguyễn Phúc Thịnh",
             size=13, align='center', space_after=4)
    add_para(doc, "MSSV: 2387700066    Lớp: 23DATA1",
             size=13, align='center', space_after=40)

    add_para(doc, "TP. Hồ Chí Minh, tháng 6 năm 2026",
             size=13, italic=True, align='center')


# ============================================================
# FRONT MATTER
# ============================================================
def build_loi_mo_dau(doc):
    add_h1(doc, "LỜI MỞ ĐẦU")
    add_para(doc,
        "Trong hơn một thập kỷ trở lại đây, xác thực sinh trắc học (biometric "
        "authentication) đã trở thành phương thức xác thực chủ đạo, được tích "
        "hợp rộng rãi trên điện thoại thông minh, hệ thống ngân hàng điện tử, "
        "kiểm soát ra vào và các dịch vụ định danh công dân điện tử. So với "
        "mật khẩu truyền thống, sinh trắc học mang lại trải nghiệm thuận tiện "
        "hơn, người dùng không cần ghi nhớ thông tin và tốc độ xác thực chỉ "
        "trong vài trăm mili-giây [1][12].",
        indent_first=1.0)
    add_para(doc,
        "Tuy nhiên, đi cùng với sự phổ biến đó là các nguy cơ bảo mật ngày "
        "càng nghiêm trọng. Các hệ thống nhận diện khuôn mặt nếu không được "
        "trang bị cơ chế kiểm tra sống (Liveness Detection hay Presentation "
        "Attack Detection) rất dễ bị qua mặt bằng ảnh in, ảnh hiển thị trên "
        "điện thoại, video phát lại hoặc các kỹ thuật injection trực tiếp ở "
        "tầng API. Khác với mật khẩu, dữ liệu sinh trắc học một khi bị lộ thì "
        "không thể thay đổi, vì vậy hậu quả của một cuộc tấn công thành công "
        "là rất lâu dài [1][9][29].",
        indent_first=1.0)
    add_para(doc,
        "Đề tài “Nghiên cứu tấn công vào hệ thống xác thực sinh trắc học” "
        "được thực hiện nhằm khảo sát một bề mặt tấn công cụ thể đối với hệ "
        "thống xác thực khuôn mặt là tấn công Injection ở tầng API, sau đó "
        "thiết kế và huấn luyện một mô hình Presentation Attack Detection "
        "dựa trên học sâu để phòng thủ. Báo cáo tập trung vào quá trình "
        "huấn luyện mô hình MobileNetV2 với kỹ thuật Transfer Learning hai "
        "pha và việc tích hợp mô hình này vào kiến trúc máy chủ thông qua "
        "lớp wrapper LivenessPredictor và hàm kiểm tra check_liveness().",
        indent_first=1.0)
    add_para(doc,
        "Mục tiêu cụ thể của đồ án bao gồm: (1) hiện thực hóa một hệ thống "
        "xác thực khuôn mặt Vulnerable bằng Flask và face_recognition để làm "
        "đối tượng nghiên cứu, (2) phân tích chi tiết kịch bản tấn công "
        "Injection bằng cách gửi ảnh đã mã hóa base64 trực tiếp vào endpoint "
        "/authenticate, (3) huấn luyện mô hình PAD MobileNetV2 với Transfer "
        "Learning hai pha trên tập dữ liệu 200 ảnh thật và 200 ảnh giả tự "
        "thu thập, (4) tích hợp mô hình vào Server Secured và (5) đánh giá "
        "định lượng tỷ lệ chặn tấn công, độ trễ phát sinh thêm cùng các chỉ "
        "số APCER/BPCER theo tiêu chuẩn ISO/IEC 30107-3.",
        indent_first=1.0)
    add_para(doc,
        "Bố cục báo cáo gồm năm chương: Chương 1 trình bày tổng quan đề tài "
        "và bối cảnh nghiên cứu; Chương 2 hệ thống cơ sở lý thuyết về sinh "
        "trắc học, PAD và các tiêu chuẩn liên quan; Chương 3 phân tích kịch "
        "bản tấn công Injection trên hệ thống Vulnerable; Chương 4 trình bày "
        "trọng tâm về huấn luyện mô hình PAD bằng MobileNetV2 và Transfer "
        "Learning hai pha cùng bước tích hợp vào Server Secured; Chương 5 "
        "tổng kết kết quả đạt được, hạn chế và đề xuất hướng phát triển.",
        indent_first=1.0)
    add_page_break(doc)


def build_loi_cam_doan(doc):
    add_h1(doc, "LỜI CAM ĐOAN")
    add_para(doc,
        "Em xin cam đoan đồ án “Nghiên cứu tấn công vào hệ thống xác thực "
        "sinh trắc học” là công trình nghiên cứu của riêng em, được thực "
        "hiện dưới sự hướng dẫn của giảng viên ThS. Đặng Thị Thạch Thảo. "
        "Các số liệu, kết quả thực nghiệm trong báo cáo là trung thực, được "
        "đo đạc trên môi trường lab do em tự xây dựng và chưa từng được công "
        "bố trong bất kỳ công trình nào khác.",
        indent_first=1.0)
    add_para(doc,
        "Các nguồn tài liệu tham khảo đều được trích dẫn rõ ràng theo định "
        "dạng [n]. Toàn bộ thí nghiệm Red Teaming chỉ được thực hiện trên hệ "
        "thống do em tự triển khai phục vụ mục đích nghiên cứu, không nhằm "
        "vào bất kỳ hệ thống thương mại của tổ chức hoặc cá nhân khác.",
        indent_first=1.0)
    add_para(doc, "TP. HCM, ngày … tháng … năm 2026",
             align='right', italic=True, space_after=2)
    add_para(doc, "Sinh viên thực hiện",
             align='right', italic=True, space_after=20)
    add_para(doc, "Nguyễn Phúc Thịnh", align='right', bold=True)
    add_page_break(doc)


def build_loi_cam_on(doc):
    add_h1(doc, "LỜI CẢM ƠN")
    add_para(doc,
        "Trước tiên, em xin gửi lời cảm ơn chân thành đến quý thầy cô Khoa "
        "Công nghệ Thông tin, Trường Đại học Công nghệ TP. HCM (HUTECH) đã "
        "trang bị cho em những kiến thức nền tảng quan trọng về an toàn thông "
        "tin, học máy và phát triển ứng dụng trong suốt thời gian học tập tại "
        "trường.",
        indent_first=1.0)
    add_para(doc,
        "Em xin gửi lời cảm ơn sâu sắc nhất đến giảng viên hướng dẫn ThS. "
        "Đặng Thị Thạch Thảo đã tận tình góp ý, định hướng và tạo điều kiện "
        "để em có thể hoàn thành đề tài. Những phản hồi kịp thời của Cô trong "
        "từng tuần báo cáo đã giúp em điều chỉnh hướng nghiên cứu phù hợp và "
        "bám sát mục tiêu đặt ra ban đầu, đặc biệt là việc tiếp cận bài toán "
        "theo tư duy Detection Engineering thay vì chỉ dừng ở mức nhận diện "
        "khuôn mặt thuần túy.",
        indent_first=1.0)
    add_para(doc,
        "Em cũng xin cảm ơn gia đình, bạn bè đã luôn động viên, hỗ trợ em "
        "trong quá trình nghiên cứu và thử nghiệm các kịch bản tấn công, "
        "phòng vệ trên hệ thống sinh trắc học.",
        indent_first=1.0)
    add_para(doc,
        "Do thời gian và kiến thức còn hạn chế, đồ án không tránh khỏi những "
        "thiếu sót. Em rất mong nhận được những góp ý quý báu từ quý thầy cô "
        "để hoàn thiện hơn trong các nghiên cứu tiếp theo. Em xin chân thành "
        "cảm ơn!",
        indent_first=1.0)
    add_page_break(doc)


def build_nhan_xet(doc):
    add_h1(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN")
    for _ in range(12):
        add_para(doc, "………………………………………………………………………………………………",
                 align='left', space_after=6, line_spacing=1.8)
    add_para(doc, "TP. HCM, ngày … tháng … năm 2026",
             align='right', italic=True, space_after=4)
    add_para(doc, "Giảng viên hướng dẫn",
             align='right', italic=True, space_after=20)
    add_para(doc, "ThS. Đặng Thị Thạch Thảo",
             align='right', bold=True)
    add_page_break(doc)


def build_muc_luc(doc):
    add_h1(doc, "MỤC LỤC")
    add_para(doc,
        "Mục lục bên dưới là TOC field tự động. Khi mở file trong Word, hãy "
        "nhấn F9 (hoặc chuột phải, chọn Update Field) để Word liệt kê đầy đủ "
        "các Heading và số trang thực tế. Đồng thời, có thể nhấn Ctrl + click "
        "vào tên mục để nhảy nhanh tới phần tương ứng.",
        size=11, italic=True, indent_first=0, space_after=8, color='555555')
    add_toc_field(doc, levels='1-3')
    add_page_break(doc)


def build_abbreviations(doc):
    add_h1(doc, "DANH MỤC CHỮ VIẾT TẮT")
    rows = [
        ["1",  "APCER",        "Attack Presentation Classification Error Rate"],
        ["2",  "BPCER",        "Bona-fide Presentation Classification Error Rate"],
        ["3",  "FAR",          "False Acceptance Rate"],
        ["4",  "FRR",          "False Rejection Rate"],
        ["5",  "EER",          "Equal Error Rate"],
        ["6",  "PAD",          "Presentation Attack Detection"],
        ["7",  "PAI",          "Presentation Attack Instrument"],
        ["8",  "TTP",          "Tactics, Techniques and Procedures"],
        ["9",  "ASR",          "Attack Success Rate"],
        ["10", "MITRE ATT&CK", "Adversarial Tactics, Techniques, and Common Knowledge"],
        ["11", "STRIDE",       "Spoofing, Tampering, Repudiation, Info Disclosure, DoS, Elevation"],
        ["12", "CNN",          "Convolutional Neural Network"],
        ["13", "MobileNetV2",  "Kiến trúc CNN nhẹ với Inverted Residual và Linear Bottleneck"],
        ["14", "ResNet",       "Residual Network"],
        ["15", "GAN",          "Generative Adversarial Network"],
        ["16", "ROC",          "Receiver Operating Characteristic"],
        ["17", "GAP",          "Global Average Pooling"],
        ["18", "BCE",          "Binary Cross Entropy"],
        ["19", "CLAHE",        "Contrast Limited Adaptive Histogram Equalization"],
        ["20", "ROI",          "Region of Interest"],
        ["21", "RGB",          "Red, Green, Blue"],
        ["22", "ISO/IEC 30107","Tiêu chuẩn quốc tế về Phát hiện tấn công trình diện"],
        ["23", "FIDO",         "Fast IDentity Online Alliance"],
        ["24", "AAL",          "Authenticator Assurance Level (NIST SP 800-63B)"],
        ["25", "OWASP",        "Open Worldwide Application Security Project"],
        ["26", "NIST",         "National Institute of Standards and Technology"],
        ["27", "API",          "Application Programming Interface"],
        ["28", "REST",         "Representational State Transfer"],
        ["29", "JSON",         "JavaScript Object Notation"],
        ["30", "JWT",          "JSON Web Token"],
        ["31", "HMAC",         "Hash-based Message Authentication Code"],
        ["32", "TLS",          "Transport Layer Security"],
        ["33", "MITM",         "Man-in-the-Middle"],
        ["34", "MFA",          "Multi-Factor Authentication"],
        ["35", "SOC",          "Security Operations Center"],
        ["36", "SIEM",         "Security Information and Event Management"],
        ["37", "SOAR",         "Security Orchestration, Automation, and Response"],
        ["38", "HSM",          "Hardware Security Module"],
        ["39", "TEE",          "Trusted Execution Environment"],
        ["40", "DB",           "Database"],
    ]
    add_table(doc, ["STT", "Viết tắt", "Nghĩa đầy đủ"], rows,
              widths=[1.5, 3.5, 11])
    add_page_break(doc)


def _add_dotleader_entry(doc, left_text, page_text):
    """Mot dong kieu muc luc: ten ben trai, dot leader, so trang ben phai."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.4
    pf.space_after = Pt(4)
    pf.left_indent = Cm(0.0)
    # Right-aligned tab with dot leader at column 15.5 cm
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_left = p.add_run(left_text)
    _report_core.set_run_font(run_left, size=12, color='000000')
    run_tab = p.add_run("\t")
    _report_core.set_run_font(run_tab, size=12, color='000000')
    run_page = p.add_run(str(page_text))
    _report_core.set_run_font(run_page, size=12, color='000000')
    return p


def build_list_of_figures(doc, figures):
    add_h1(doc, "DANH MỤC HÌNH ẢNH")
    if not figures:
        _add_dotleader_entry(doc,
            "(chưa có hình nào, sẽ được cập nhật sau khi build chương)",
            "—")
        add_page_break(doc)
        return
    for item in figures:
        if len(item) == 3:
            num, title, page = item
        else:
            num, title = item
            page = "(chưa có)"
        _add_dotleader_entry(doc, f"Hình {num}. {title}", page)
    add_page_break(doc)


def build_list_of_tables(doc, tables):
    add_h1(doc, "DANH MỤC BẢNG BIỂU")
    if not tables:
        _add_dotleader_entry(doc,
            "(chưa có bảng nào, sẽ được cập nhật sau khi build chương)",
            "—")
        add_page_break(doc)
        return
    for item in tables:
        if len(item) == 3:
            num, title, page = item
        else:
            num, title = item
            page = "(chưa có)"
        _add_dotleader_entry(doc, f"Bảng {num}. {title}", page)
    add_page_break(doc)


# ============================================================
# DRY-RUN PASS
# ============================================================
def _dry_run_collect():
    dry = Document()
    setup_doc_defaults(dry)
    build_chapter1(dry)
    build_chapter2(dry)
    build_chapter3(dry)
    build_chapter4(dry)
    build_chapter5(dry)
    figures = list(_report_core.FIGURES_LOG)
    tables = list(_report_core.TABLES_LOG)
    _report_core.FIGURES_LOG.clear()
    _report_core.TABLES_LOG.clear()
    _report_core.PAGE_TRACKER["counter"] = 0
    return figures, tables


# ============================================================
# MAIN
# ============================================================
def main():
    print("[1/3] Dry-run: thu thap Figures/Tables...")
    figures, tables = _dry_run_collect()
    print(f"     -> {len(figures)} hinh, {len(tables)} bang.")

    print("[2/3] Build doc that...")
    doc = Document()
    setup_doc_defaults(doc)

    # ----- Section 1: Cover (no page numbering) -----
    build_cover(doc)

    # ----- Section 2: Front matter (upperRoman) -----
    new_section(doc, fmt='upperRoman', start=1)
    build_loi_mo_dau(doc)
    build_loi_cam_doan(doc)
    build_loi_cam_on(doc)
    build_nhan_xet(doc)
    build_muc_luc(doc)
    build_abbreviations(doc)
    build_list_of_figures(doc, figures)
    build_list_of_tables(doc, tables)

    # ----- Section 3: Chapters (decimal, restart 1) -----
    new_section(doc, fmt='decimal', start=1)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_chapter4(doc)
    build_chapter5(doc)
    build_references(doc)

    print(f"[3/3] Save -> {OUT_FINAL}")
    doc.save(OUT_FINAL)
    size_kb = os.path.getsize(OUT_FINAL) / 1024
    paras = _report_core.PAGE_TRACKER["counter"]
    pages = max(1, paras // 35 + 1)
    print(f"     -> OK, {size_kb:.1f} KB")
    print(f"     -> paragraph counter = {paras}, uoc tinh ~ {pages} trang")


if __name__ == "__main__":
    main()
